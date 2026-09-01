from __future__ import annotations

import json
import hashlib
import tempfile
import unittest
from pathlib import Path
from types import SimpleNamespace
from unittest.mock import patch

from PIL import Image
from docx import Document

from server.pdf_to_docx_core import positioned, reporting, runtime


class RuntimeAndReportingTests(unittest.TestCase):
    def test_find_tesseract_prefers_path_lookup(self) -> None:
        with patch.object(runtime.shutil, "which", return_value="/custom/tesseract"):
            self.assertEqual(runtime.find_tesseract(), "/custom/tesseract")

    def test_system_language_detection_parses_tesseract_output(self) -> None:
        result = SimpleNamespace(
            stdout="List of available languages (3):\neng\nrus\ntgk\n", stderr=""
        )
        with patch.object(runtime.subprocess, "run", return_value=result):
            self.assertEqual(
                runtime.system_tesseract_languages("tesseract"), {"eng", "rus", "tgk"}
            )

    def test_system_language_detection_is_safe_when_process_fails(self) -> None:
        with patch.object(runtime.subprocess, "run", side_effect=OSError("missing")):
            self.assertEqual(runtime.system_tesseract_languages("missing"), set())

    def test_local_language_detection_only_includes_traineddata(self) -> None:
        with (
            tempfile.TemporaryDirectory() as tmp,
            patch.object(runtime, "LOCAL_TESSDATA_DIR", tmp),
        ):
            Path(tmp, "eng.traineddata").write_bytes(b"model")
            Path(tmp, "README.txt").write_text("ignore", encoding="utf-8")
            self.assertEqual(runtime.local_tesseract_languages(), {"eng"})

    def test_tessdata_source_is_revision_pinned_with_complete_sha256_map(self) -> None:
        from server import pdf_to_docx_config as config

        self.assertRegex(config.TESSDATA_FAST_REVISION, r"^[0-9a-f]{40}$")
        self.assertIn(config.TESSDATA_FAST_REVISION, config.TESSDATA_FAST_BASE)
        self.assertNotIn("/main", config.TESSDATA_FAST_BASE)
        self.assertEqual(set(config.TESSDATA_FAST_SHA256), config.DOWNLOADABLE_TESSDATA)
        for digest in config.TESSDATA_FAST_SHA256.values():
            self.assertRegex(digest, r"^[0-9a-f]{64}$")

    def test_tessdata_download_installs_only_matching_content(self) -> None:
        payload = b"verified-model" * 200
        expected = hashlib.sha256(payload).hexdigest()
        report: dict[str, list[str]] = {"notes": [], "warnings": []}

        def download(_url: str, destination: str) -> None:
            Path(destination).write_bytes(payload)

        with (
            tempfile.TemporaryDirectory() as tmp,
            patch.object(runtime, "LOCAL_TESSDATA_DIR", tmp),
            patch.object(runtime, "TESSDATA_FAST_SHA256", {"fas": expected}),
            patch.object(runtime, "installed_tesseract_languages", return_value=set()),
            patch.object(runtime.urllib.request, "urlretrieve", side_effect=download),
        ):
            self.assertTrue(runtime.ensure_project_tessdata("fas", report))
            self.assertEqual(Path(tmp, "fas.traineddata").read_bytes(), payload)
        self.assertIn("Downloaded OCR language data: fas.", report["notes"])

    def test_tessdata_download_rejects_and_removes_checksum_mismatch(self) -> None:
        report: dict[str, list[str]] = {"notes": [], "warnings": []}

        def download(_url: str, destination: str) -> None:
            Path(destination).write_bytes(b"untrusted-model" * 200)

        with (
            tempfile.TemporaryDirectory() as tmp,
            patch.object(runtime, "LOCAL_TESSDATA_DIR", tmp),
            patch.object(runtime, "TESSDATA_FAST_SHA256", {"fas": "0" * 64}),
            patch.object(runtime, "installed_tesseract_languages", return_value=set()),
            patch.object(runtime.urllib.request, "urlretrieve", side_effect=download),
        ):
            self.assertFalse(runtime.ensure_project_tessdata("fas", report))
            self.assertFalse(Path(tmp, "fas.traineddata").exists())
            self.assertFalse(Path(tmp, "fas.traineddata.download").exists())
        self.assertIn("checksum", report["warnings"][0])

    def test_tessdata_valid_cached_model_skips_network(self) -> None:
        payload = b"cached-model" * 200
        expected = hashlib.sha256(payload).hexdigest()
        report: dict[str, list[str]] = {"notes": [], "warnings": []}
        with (
            tempfile.TemporaryDirectory() as tmp,
            patch.object(runtime, "LOCAL_TESSDATA_DIR", tmp),
            patch.object(runtime, "TESSDATA_FAST_SHA256", {"fas": expected}),
            patch.object(runtime, "installed_tesseract_languages", return_value=set()),
            patch.object(runtime.urllib.request, "urlretrieve") as download,
        ):
            Path(tmp, "fas.traineddata").write_bytes(payload)
            self.assertTrue(runtime.ensure_project_tessdata("fas", report))
            download.assert_not_called()

    def test_resolve_language_keeps_installed_manual_languages_and_warns_for_missing(
        self,
    ) -> None:
        report: dict[str, object] = {"notes": [], "warnings": []}
        with (
            patch.object(runtime, "find_tesseract", return_value="tesseract"),
            patch.object(runtime, "FAST_HOSTED_OCR", True),
            patch.object(
                runtime, "installed_tesseract_languages", return_value={"eng", "rus"}
            ),
            patch.object(runtime, "local_tesseract_languages", return_value=set()),
        ):
            language = runtime.resolve_ocr_language("eng+tgk", report)
        self.assertEqual(language, "eng")
        self.assertIn(
            "Missing Tesseract language data: tgk", str(report["warnings"][0])
        )

    def test_resolve_auto_language_records_limited_install_warning(self) -> None:
        report: dict[str, object] = {"notes": [], "warnings": []}
        with (
            patch.object(runtime, "find_tesseract", return_value="tesseract"),
            patch.object(runtime, "FAST_HOSTED_OCR", True),
            patch.object(
                runtime, "installed_tesseract_languages", return_value={"eng"}
            ),
            patch.object(runtime, "local_tesseract_languages", return_value=set()),
        ):
            self.assertEqual(runtime.resolve_ocr_language("auto", report), "eng")
        self.assertTrue(report["warnings"])
        self.assertEqual(report["ocrLanguage"], "eng")

    def test_inspect_pdf_identifies_native_text_and_has_text_layer(self) -> None:
        import fitz

        with tempfile.TemporaryDirectory() as tmp:
            source = Path(tmp) / "native.pdf"
            pdf = fitz.open()
            page = pdf.new_page()
            page.insert_text(
                (50, 50), "This native PDF has more than twenty-five characters."
            )
            pdf.save(source)
            pdf.close()
            info = runtime.inspect_pdf(str(source))
            self.assertTrue(runtime.has_text_layer(str(source)))
        self.assertEqual(info["pdfType"], "digital")
        self.assertEqual(info["pages"], 1)
        self.assertGreater(info["textCharacters"], 25)

    def test_image_backed_heuristics_use_page_coverage_and_text_density(self) -> None:
        info = {
            "pages": 2,
            "pageDetails": [
                {"maxImageCoverage": 0.9, "textCharacters": 120},
                {"maxImageCoverage": 0.7, "textCharacters": 500},
            ],
        }
        self.assertTrue(positioned.image_backed_text_layer_likely(info))
        self.assertTrue(positioned.image_backed_text_layer_needs_ocr(info))
        self.assertFalse(
            positioned.image_backed_text_layer_likely({"pages": 0, "pageDetails": []})
        )

    def test_ensure_output_rejects_missing_and_empty_files(self) -> None:
        with tempfile.TemporaryDirectory() as tmp:
            missing = Path(tmp) / "missing.docx"
            with self.assertRaisesRegex(SystemExit, "no output"):
                reporting.ensure_output(str(missing))
            missing.write_bytes(b"")
            with self.assertRaises(SystemExit):
                reporting.ensure_output(str(missing))
            missing.write_bytes(b"ok")
            reporting.ensure_output(str(missing))

    def test_docx_stats_count_text_tables_and_images(self) -> None:
        with tempfile.TemporaryDirectory() as tmp:
            image_path = Path(tmp) / "pixel.png"
            output = Path(tmp) / "document.docx"
            Image.new("RGB", (8, 8), "blue").save(image_path)
            document = Document()
            document.add_paragraph("Editable FileMint text")
            table = document.add_table(rows=1, cols=1)
            table.cell(0, 0).text = "Table value"
            document.add_picture(str(image_path))
            document.save(output)
            stats = reporting.docx_output_stats(str(output))
        self.assertGreaterEqual(stats["outputTextRuns"], 2)
        self.assertGreaterEqual(stats["outputEditableCharacters"], 20)
        self.assertGreaterEqual(stats["outputTables"], 1)
        self.assertEqual(stats["outputImages"], 1)

    def test_docx_stats_return_zeroes_for_invalid_archives(self) -> None:
        with tempfile.TemporaryDirectory() as tmp:
            invalid = Path(tmp) / "invalid.docx"
            invalid.write_text("not a zip", encoding="utf-8")
            self.assertEqual(
                reporting.docx_output_stats(str(invalid)),
                {
                    "outputTextRuns": 0,
                    "outputEditableCharacters": 0,
                    "outputImages": 0,
                    "outputTables": 0,
                },
            )

    def test_merge_output_stats_updates_report_in_place(self) -> None:
        with tempfile.TemporaryDirectory() as tmp:
            output = Path(tmp) / "document.docx"
            document = Document()
            document.add_paragraph("text")
            document.save(output)
            report: dict[str, object] = {"engine": "filemint"}
            stats = reporting.merge_output_stats(report, str(output))
        self.assertEqual(report["outputTextRuns"], stats["outputTextRuns"])
        self.assertEqual(report["engine"], "filemint")

    def test_write_report_serializes_unicode_as_portable_json(self) -> None:
        with tempfile.TemporaryDirectory() as tmp:
            output = Path(tmp) / "report.json"
            reporting.write_report(str(output), {"language": "Тоҷикӣ"})
            raw = output.read_text(encoding="utf-8")
            parsed = json.loads(raw)
        self.assertEqual(parsed["language"], "Тоҷикӣ")
        self.assertIn("\\u", raw)

    def test_repair_uses_exact_visual_path_when_available(self) -> None:
        report: dict[str, object] = {"warnings": []}
        with (
            patch.object(reporting, "resolve_ocr_language", return_value="eng"),
            patch.object(reporting, "ocr_to_docx_exact_visual") as exact,
            patch.object(reporting, "ocr_to_docx_layout") as layout,
        ):
            reporting.repair_empty_editable_output(
                "in.pdf", "out.docx", "eng", True, report, "high", True
            )
        exact.assert_called_once()
        layout.assert_not_called()
        self.assertEqual(report["resolvedMode"], "ocr-repair-editable-visual")

    def test_repair_falls_back_to_text_flow_after_exact_visual_failure(self) -> None:
        report: dict[str, object] = {"warnings": []}
        with (
            patch.object(reporting, "resolve_ocr_language", return_value="eng"),
            patch.object(
                reporting,
                "ocr_to_docx_exact_visual",
                side_effect=RuntimeError("visual failed"),
            ),
            patch.object(reporting, "ocr_to_docx_layout") as layout,
        ):
            reporting.repair_empty_editable_output(
                "in.pdf", "out.docx", "eng", True, report, "high", False
            )
        layout.assert_called_once()
        self.assertEqual(report["resolvedMode"], "ocr-repair-text-flow")
        self.assertIn("visual failed", str(report["warnings"][-1]))


if __name__ == "__main__":
    unittest.main()
