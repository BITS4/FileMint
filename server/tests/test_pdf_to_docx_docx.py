"""Branch coverage for reusable DOCX text and table formatting helpers."""

from __future__ import annotations

import unittest
from unittest.mock import patch

from docx import Document
from docx.oxml.ns import qn

from server.pdf_to_docx_core import docx
from server.pdf_to_docx_core.models import LineBox, WordBox


def line(
    text: str,
    *,
    left: float = 20,
    top: float = 20,
    width: float = 100,
    height: float = 12,
    conf: float = 90,
    segments: list[tuple[float, float, str]] | None = None,
) -> LineBox:
    word = WordBox(text, left, top, width, height, conf, 1, 1, 1)
    return LineBox(
        text=text,
        words=[word],
        left=left,
        top=top,
        width=width,
        height=height,
        conf=conf,
        page_width_px=600,
        page_height_px=800,
        page_width_pt=300,
        page_height_pt=400,
        segments=segments if segments is not None else [(left, left + width, text)],
    )


class DocxFormattingTests(unittest.TestCase):
    def test_rtl_detection_and_bidi_marker_are_idempotent(self) -> None:
        self.assertTrue(docx.contains_rtl("سلام"))
        self.assertFalse(docx.contains_rtl("FileMint 42"))

        paragraph = Document().add_paragraph()
        docx.set_paragraph_bidi(paragraph)
        docx.set_paragraph_bidi(paragraph)
        self.assertEqual(
            len(paragraph._p.xpath("./w:pPr/w:bidi")),
            1,
        )

    def test_run_font_creates_then_reuses_multiscript_font_properties(self) -> None:
        run = Document().add_paragraph().add_run("Mixed")
        docx.set_run_font(run, "Calibri")
        docx.set_run_font(run, "Arial")

        fonts = run._element.get_or_add_rPr().rFonts
        self.assertIsNotNone(fonts)
        for key in ("w:ascii", "w:hAnsi", "w:eastAsia", "w:cs"):
            self.assertEqual(fonts.get(qn(key)), "Arial")

    def test_cell_text_sets_direction_font_and_bounded_size(self) -> None:
        table = Document().add_table(rows=1, cols=2)
        docx.set_cell_text(table.cell(0, 0), "مرحبا", 2, True)
        docx.set_cell_text(table.cell(0, 1), "Large", 99, False)

        self.assertEqual(table.cell(0, 0).text, "مرحبا")
        self.assertEqual(
            table.cell(0, 0).paragraphs[0]._p.xpath("./w:pPr/w:bidi")[0].tag,
            qn("w:bidi"),
        )
        self.assertAlmostEqual(table.cell(0, 0).paragraphs[0].runs[-1].font.size.pt, 6)
        self.assertAlmostEqual(table.cell(0, 1).paragraphs[0].runs[-1].font.size.pt, 22)

        with patch("docx.shared.Pt", side_effect=RuntimeError("font unavailable")):
            docx.set_cell_text(table.cell(0, 1), "Fallback", 10, False)
        self.assertEqual(table.cell(0, 1).text, "Fallback")

    def test_column_clustering_handles_empty_nearby_and_separated_positions(
        self,
    ) -> None:
        self.assertEqual(docx.cluster_columns([], 3), [])
        lines = [
            line(
                "row 1",
                segments=[(10, 50, "A"), (100, 140, "B"), (220, 250, "C")],
            ),
            line(
                "row 2",
                segments=[(14, 50, "D"), (105, 140, "E"), (300, 330, "F")],
            ),
        ]
        self.assertEqual(docx.cluster_columns(lines, 3), [12.0, 102.5, 220])

    def test_table_run_detection_splits_on_noise_geometry_and_column_changes(
        self,
    ) -> None:
        two = [(10, 40, "A"), (80, 120, "B")]
        three = [*two, (150, 180, "C")]
        lines = [
            line("single", top=0, segments=[(10, 40, "single")]),
            line("r1", top=20, segments=two),
            line("r2", top=35, segments=three),
            line("break", top=50, segments=[(10, 40, "break")]),
            line("r3", top=100, segments=two),
            line("r4", top=300, segments=two),
            line("r5", top=315, segments=two),
        ]

        runs = docx.table_runs(lines)
        self.assertEqual(
            [[item.text for item in run] for run in runs], [["r1", "r2"], ["r4", "r5"]]
        )

    def test_add_table_builds_missing_cells_and_updates_report(self) -> None:
        source = [
            line(
                "first",
                height=18,
                segments=[(10, 70, "Name"), (120, 170, "Value"), (220, 260, "Extra")],
            ),
            line(
                "second",
                top=45,
                segments=[(10, 70, "Alice"), (120, 170, "42")],
            ),
        ]
        document = Document()
        report: dict[str, object] = {"tablesDetected": "2"}
        docx.add_table(document, source, report)

        table = document.tables[0]
        self.assertEqual(len(table.rows), 2)
        self.assertEqual(len(table.columns), 3)
        self.assertEqual(
            [cell.text for cell in table.rows[1].cells], ["Alice", "42", ""]
        )
        self.assertEqual(report["tablesDetected"], 3)

    def test_paragraph_layout_handles_ltr_rtl_and_previous_line_spacing(self) -> None:
        document = Document()
        latin = line("English paragraph", left=5, top=10, height=10)
        rtl = line("متن فارسی", left=100, top=50, height=20)
        docx.add_paragraph_from_line(document, latin, None)
        docx.add_paragraph_from_line(document, rtl, 20)

        self.assertEqual(document.paragraphs[0].text, latin.text)
        self.assertEqual(document.paragraphs[0].paragraph_format.left_indent.pt, 0)
        self.assertEqual(
            document.paragraphs[1]._p.xpath("./w:pPr/w:bidi")[0].tag, qn("w:bidi")
        )
        self.assertGreater(document.paragraphs[1].paragraph_format.space_before.pt, 0)

    def test_ocr_font_metrics_languages_and_text_quality_thresholds(self) -> None:
        self.assertEqual(docx.ocr_font_size(line("", height=2)), 5.0)
        fitted = docx.ocr_font_size(line("A very long text 123", width=20, height=30))
        self.assertGreaterEqual(fitted, 4.5)
        self.assertLessEqual(fitted, 20)

        self.assertEqual(
            docx.ocr_font_attrs("漢字"), ("Times New Roman", "SimSun", "Arial")
        )
        self.assertEqual(docx.ocr_font_attrs("سلام"), ("Arial", "Arial", "Arial"))
        self.assertEqual(
            docx.ocr_font_attrs("Latin"),
            ("Times New Roman", "Times New Roman", "Arial"),
        )
        self.assertEqual(docx.editable_confidence_threshold("eng+rus"), 42.0)
        self.assertEqual(docx.editable_confidence_threshold("eng"), 72.0)
        self.assertEqual(docx.premium_confidence_threshold("fas"), 18.0)
        self.assertEqual(docx.premium_confidence_threshold("eng"), 38.0)

        self.assertFalse(docx.line_is_confident(line("good", conf=30)))
        self.assertFalse(docx.line_is_confident(line("!", conf=99)))
        self.assertTrue(docx.line_is_confident(line("A1", conf=99)))
        self.assertFalse(docx.line_text_signal(line("!!")))
        self.assertTrue(docx.line_text_signal(line("42")))
        self.assertTrue(docx.line_text_signal(line("AB!")))


if __name__ == "__main__":
    unittest.main()
