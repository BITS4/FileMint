from __future__ import annotations

import tempfile
import unittest
from pathlib import Path

from PIL import Image
from docx import Document

from server.pdf_to_docx_core.models import LineBox, VisualFragment, VisualRule, WordBox
from server.pdf_to_docx_core.native import PdfLine, PdfSpan
from server.pdf_to_docx_core.positioned import (
    append_exact_visual_page,
    append_linebox_flow_paragraph,
    append_positioned_page,
    native_pdf_line_boxes,
    pdf_line_text_with_gaps,
)


def make_line(
    text: str = "Editable text", *, left: float = 20, top: float = 30
) -> LineBox:
    word = WordBox(text, left, top, 100, 18, 95, 1, 1, 1)
    return LineBox(
        text,
        [word],
        left,
        top,
        100,
        18,
        95,
        400,
        600,
        200,
        300,
        [(left, left + 100, text)],
    )


class PositionedDocxTests(unittest.TestCase):
    def test_pdf_line_text_infers_spaces_from_span_geometry(self) -> None:
        line = PdfLine(
            [
                PdfSpan("File", "Times", 12, 0, 0, (0, 0, 20, 12)),
                PdfSpan("Mint", "Times", 12, 0, 0, (35, 0, 60, 12)),
            ],
            0,
            0,
            60,
            12,
        )
        self.assertEqual(pdf_line_text_with_gaps(line), "File    Mint")

    def test_native_pdf_line_boxes_scale_native_coordinates(self) -> None:
        import fitz

        pdf = fitz.open()
        page = pdf.new_page(width=200, height=300)
        page.insert_text((20, 50), "Native line geometry")
        try:
            lines = native_pdf_line_boxes(page, 400, 600)
        finally:
            pdf.close()
        self.assertEqual(len(lines), 1)
        self.assertEqual(lines[0].page_width_px, 400)
        self.assertEqual(lines[0].page_width_pt, 200)
        self.assertGreater(lines[0].width, 100)

    def test_flow_paragraph_normalizes_text_and_returns_character_count(self) -> None:
        document = Document()
        line = make_line("  Editable   text  ")
        emitted = append_linebox_flow_paragraph(document, line, None, 20, 20)
        self.assertEqual(document.paragraphs[-1].text, "Editable text")
        self.assertEqual(emitted, len("Editable text"))

    def test_flow_paragraph_skips_empty_text(self) -> None:
        document = Document()
        self.assertEqual(
            append_linebox_flow_paragraph(document, make_line("   "), None, 20, 20), 0
        )
        self.assertEqual(len(document.paragraphs), 0)

    def test_exact_visual_page_embeds_background_and_hidden_editable_text(self) -> None:
        with tempfile.TemporaryDirectory() as tmp:
            image_path = Path(tmp) / "page.png"
            Image.new("RGB", (200, 300), "white").save(image_path)
            document = Document()
            append_exact_visual_page(
                document, 0, str(image_path), [make_line()], 200, 300, hidden_text=True
            )
            xml = document.paragraphs[-1]._p.xml
        self.assertIn("FileMintText0_0", xml)
        self.assertIn("w:vanish", xml)
        self.assertIn("Editable text", xml)

    def test_positioned_page_embeds_rules_fragments_and_editable_text(self) -> None:
        with tempfile.TemporaryDirectory() as tmp:
            fragment_path = Path(tmp) / "fragment.png"
            Image.new("RGBA", (20, 20), (255, 0, 0, 255)).save(fragment_path)
            document = Document()
            append_positioned_page(
                document,
                2,
                [make_line()],
                200,
                300,
                [VisualFragment(str(fragment_path), 10, 10, 20, 20)],
                [VisualRule(5, 100, 120, 1, "#123456")],
            )
            xml = document.paragraphs[-1]._p.xml
        self.assertIn("FileMintRule2_0", xml)
        self.assertIn("FileMintFragment2_0", xml)
        self.assertIn("FileMintText2_0", xml)
        self.assertIn("#123456", xml)


if __name__ == "__main__":
    unittest.main()
