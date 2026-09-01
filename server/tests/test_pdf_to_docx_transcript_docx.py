"""Branch coverage for scanned-transcript Word table reconstruction."""

from __future__ import annotations

import tempfile
import unittest
from pathlib import Path
from unittest.mock import patch

from docx import Document
from PIL import Image, ImageDraw

from server.pdf_to_docx_core import transcript_docx
from server.pdf_to_docx_core.models import LineBox, WordBox


def line(
    text: str,
    *,
    left: float,
    top: float,
    conf: float = 90,
    width: float = 45,
    height: float = 8,
) -> LineBox:
    word = WordBox(text, left, top, width, height, conf, 1, 1, 1)
    return LineBox(
        text=text,
        words=[word],
        left=left,
        top=top,
        width=width,
        height=height,
        conf=conf,
        page_width_px=200,
        page_height_px=300,
        page_width_pt=400,
        page_height_pt=600,
        segments=[(left, left + width, text)],
    )


def source_lines() -> list[LineBox]:
    return [
        line("School", left=10, top=12),
        line("Transcript", left=85, top=16),
        line("Year 2026", left=150, top=20),
        line("ignored", left=20, top=25, conf=20),
        line("Mathematics", left=55, top=112),
        line("A", left=120, top=112),
        line("Physics", left=55, top=130),
        line("B", left=150, top=130),
    ]


class TranscriptDocxTests(unittest.TestCase):
    def create_scan(self, directory: str) -> str:
        path = Path(directory, "transcript.png")
        image = Image.new("RGB", (200, 300), "white")
        draw = ImageDraw.Draw(image)
        draw.rectangle((80, 260, 120, 285), fill=(20, 80, 190))
        image.save(path)
        return str(path)

    def test_measured_grid_rebuilds_table_geometry_and_preserves_seal(self) -> None:
        geometry = {
            "xPositionsPx": [20, 45, 105, 130, 160, 190],
            "yPositionsPx": [90 + index * 7 for index in range(26)],
            "tableWidthPt": 340,
            "tableLeftPt": 38,
            "colWidthsPt": [34, 150, 52, 52, 52],
            "rowHeightsPt": [18 + (index % 2) for index in range(25)],
            "preTableSpacerPt": 3,
        }
        report: dict[str, list[str]] = {"notes": []}
        document = Document()

        with (
            tempfile.TemporaryDirectory() as tmp,
            patch.object(
                transcript_docx,
                "detect_transcript_grid_geometry",
                return_value=geometry,
            ),
        ):
            transcript_docx.build_scanned_table_page(
                document,
                self.create_scan(tmp),
                source_lines(),
                400,
                600,
                report,
            )

        self.assertEqual(len(document.tables), 3)
        academic_table = document.tables[-1]
        self.assertEqual(len(academic_table.rows), 24)
        self.assertEqual(len(academic_table.columns), 5)
        self.assertEqual(academic_table.cell(1, 0).text, "1")
        self.assertTrue(
            any("geometry was measured" in note for note in report["notes"])
        )
        self.assertTrue(any("real Word table" in note for note in report["notes"]))
        self.assertTrue(document.inline_shapes or document.part.related_parts)

    def test_fallback_geometry_uses_detected_centers_when_available(self) -> None:
        report: dict[str, list[str]] = {"notes": []}
        document = Document()
        centers = [108.0, 132.0, 156.0]

        with (
            tempfile.TemporaryDirectory() as tmp,
            patch.object(
                transcript_docx, "detect_transcript_grid_geometry", return_value=None
            ),
            patch.object(
                transcript_docx, "transcript_row_centers", return_value=centers
            ),
        ):
            transcript_docx.build_scanned_table_page(
                document,
                self.create_scan(tmp),
                source_lines(),
                400,
                600,
                report,
            )

        table = document.tables[-1]
        self.assertEqual(len(table.rows), 4)
        self.assertEqual(
            [table.cell(index, 0).text for index in range(1, 4)], ["1", "2", "3"]
        )
        self.assertFalse(
            any("geometry was measured" in note for note in report["notes"])
        )

    def test_fallback_geometry_synthesizes_rows_when_ocr_has_no_centers(self) -> None:
        report: dict[str, list[str]] = {"notes": []}
        document = Document()

        with (
            tempfile.TemporaryDirectory() as tmp,
            patch.object(
                transcript_docx, "detect_transcript_grid_geometry", return_value=None
            ),
            patch.object(transcript_docx, "transcript_row_centers", return_value=[]),
        ):
            transcript_docx.build_scanned_table_page(
                document,
                self.create_scan(tmp),
                source_lines(),
                400,
                600,
                report,
            )

        self.assertEqual(len(document.tables[-1].rows), 24)
        self.assertEqual(document.tables[-1].cell(23, 0).text, "23")


if __name__ == "__main__":
    unittest.main()
