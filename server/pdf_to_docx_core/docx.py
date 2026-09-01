"""DOCX text, font, table, and paragraph formatting helpers."""

from __future__ import annotations

import re
from statistics import median
from typing import Any

from .models import LineBox
from .ocr import CJK_RANGE, script_counts


def contains_rtl(text: str) -> bool:
    return bool(re.search(r"[\u0590-\u08ff]", text))


def set_paragraph_bidi(paragraph: Any) -> None:
    from docx.oxml import OxmlElement

    p_pr = paragraph._p.get_or_add_pPr()
    if p_pr.find("./w:bidi", paragraph._p.nsmap) is None:
        p_pr.append(OxmlElement("w:bidi"))


def set_run_font(run: Any, name: str = "Arial") -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    run.font.name = name
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.append(r_fonts)
    for key in ("w:ascii", "w:hAnsi", "w:eastAsia", "w:cs"):
        r_fonts.set(qn(key), name)


def set_cell_text(cell: Any, text: str, font_size: float, rtl: bool) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    if rtl:
        set_paragraph_bidi(p)
    run = p.add_run(text)
    set_run_font(run)
    try:
        from docx.shared import Pt

        run.font.size = Pt(max(6.0, min(22.0, font_size)))
    except Exception:
        pass


def cluster_columns(lines: list[LineBox], max_cols: int) -> list[float]:
    lefts: list[float] = []
    for line in lines:
        for left, _right, _text in line.segments:
            lefts.append(left)
    if not lefts:
        return []
    lefts.sort()
    clusters: list[list[float]] = []
    for x in lefts:
        if not clusters or abs(median(clusters[-1]) - x) > 45:
            clusters.append([x])
        else:
            clusters[-1].append(x)
    centers = [median(c) for c in clusters]
    return centers[:max_cols]


def table_runs(lines: list[LineBox]) -> list[list[LineBox]]:
    runs: list[list[LineBox]] = []
    current: list[LineBox] = []
    last: LineBox | None = None
    for line in lines:
        is_candidate = len(line.segments) >= 2
        if not is_candidate:
            if len(current) >= 2:
                runs.append(current)
            current = []
            last = None
            continue

        if last is None:
            current = [line]
        else:
            vertical_gap = line.top - (last.top + last.height)
            similar_cols = abs(len(line.segments) - len(last.segments)) <= 1
            close = vertical_gap <= max(80, median([last.height, line.height]) * 3.5)
            if similar_cols and close:
                current.append(line)
            else:
                if len(current) >= 2:
                    runs.append(current)
                current = [line]
        last = line

    if len(current) >= 2:
        runs.append(current)
    return runs


def add_table(doc: Any, lines: list[LineBox], report: dict[str, Any]) -> None:
    from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
    from docx.shared import Pt

    max_cols = max(len(l.segments) for l in lines)
    table = doc.add_table(rows=len(lines), cols=max_cols)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    columns = cluster_columns(lines, max_cols)
    page_width_pt = lines[0].page_width_pt
    scale = page_width_pt / max(1.0, lines[0].page_width_px)
    usable_width = max(120.0, page_width_pt - 72.0)

    col_widths: list[float] = []
    for i in range(max_cols):
        left = (
            columns[i] * scale
            if i < len(columns)
            else (36 + (usable_width / max_cols) * i)
        )
        if i + 1 < len(columns):
            right = columns[i + 1] * scale
        else:
            right = page_width_pt - 36
        col_widths.append(max(36.0, right - left))

    for r_idx, line in enumerate(lines):
        row = table.rows[r_idx]
        try:
            row.height = Pt(max(14.0, line.height * scale * 1.35))
        except Exception:
            pass
        for c_idx in range(max_cols):
            cell = row.cells[c_idx]
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            try:
                cell.width = Pt(col_widths[c_idx])
            except Exception:
                pass
            text = line.segments[c_idx][2] if c_idx < len(line.segments) else ""
            set_cell_text(
                cell, text, max(8.0, line.height * scale * 0.95), contains_rtl(text)
            )

    report["tablesDetected"] = int(report.get("tablesDetected", 0)) + 1


def add_paragraph_from_line(doc: Any, line: LineBox, prev_bottom: float | None) -> None:
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt

    scale_x = line.page_width_pt / max(1.0, line.page_width_px)
    scale_y = line.page_height_pt / max(1.0, line.page_height_px)
    p = doc.add_paragraph()
    if contains_rtl(line.text):
        set_paragraph_bidi(p)
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    else:
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    fmt = p.paragraph_format
    fmt.left_indent = Pt(max(0.0, line.left * scale_x - 36.0))
    if prev_bottom is not None:
        gap = max(0.0, line.top - prev_bottom)
        fmt.space_before = Pt(min(36.0, gap * scale_y))
    fmt.space_after = Pt(0)
    fmt.line_spacing = 1.0

    run = p.add_run(line.text)
    set_run_font(run)
    run.font.size = Pt(max(6.0, min(22.0, line.height * scale_y * 0.95)))


VML_NS = (
    'xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" '
    'xmlns:v="urn:schemas-microsoft-com:vml" '
    'xmlns:o="urn:schemas-microsoft-com:office:office" '
    'xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships"'
)


def ocr_font_size(line: LineBox) -> float:
    scale_y = line.page_height_pt / max(1.0, line.page_height_px)
    scale_x = line.page_width_pt / max(1.0, line.page_width_px)
    height_size = line.height * scale_y * 0.68
    counts = script_counts(line.text)
    spaces = line.text.count(" ")
    punctuation = sum(1 for ch in line.text if not ch.isalnum() and not ch.isspace())
    units = (
        counts["latin"] * 0.48
        + counts["digits"] * 0.48
        + counts["cyrillic"] * 0.52
        + counts["rtl"] * 0.56
        + counts["cjk"] * 0.96
        + spaces * 0.26
        + punctuation * 0.30
    )
    if units <= 0:
        return max(5.0, min(20.0, height_size))
    fit_size = (line.width * scale_x * 1.04 + 8.0) / units
    return max(4.5, min(20.0, height_size, fit_size))


def ocr_font_attrs(text: str) -> tuple[str, str, str]:
    if re.search(rf"[{CJK_RANGE}]", text):
        return ("Times New Roman", "SimSun", "Arial")
    if contains_rtl(text):
        return ("Arial", "Arial", "Arial")
    return ("Times New Roman", "Times New Roman", "Arial")


def editable_confidence_threshold(lang: str) -> float:
    parts = set(p for p in re.split(r"[,+\s]+", lang) if p)
    if parts & {"rus", "tgk", "fas", "ara", "chi_sim", "chi_tra"}:
        return 42.0
    return 72.0


def premium_confidence_threshold(lang: str) -> float:
    parts = set(p for p in re.split(r"[,+\s]+", lang) if p)
    if parts & {"rus", "tgk", "fas", "ara", "chi_sim", "chi_tra"}:
        return 18.0
    return 38.0


def line_is_confident(line: LineBox, min_conf: float = 72.0) -> bool:
    if line.conf < min_conf:
        return False
    alnum = sum(1 for ch in line.text if ch.isalnum())
    return alnum >= 2


def line_text_signal(line: LineBox) -> bool:
    alnum = sum(1 for ch in line.text if ch.isalnum())
    letters = sum(1 for ch in line.text if ch.isalpha())
    return alnum >= 2 and (letters >= 1 or any(ch.isdigit() for ch in line.text))
