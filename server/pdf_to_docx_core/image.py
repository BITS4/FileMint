"""Layout OCR and image-only DOCX conversion modes."""

from __future__ import annotations

import os
import shutil
import tempfile
from typing import Any

from .config import quality_dpi
from .docx import add_paragraph_from_line, add_table, table_runs
from .ocr import parse_tsv, run_tesseract_tsv


def ocr_to_docx_layout(
    src: str,
    dst: str,
    lang: str,
    table_detection: bool,
    report: dict[str, Any],
    quality: str = "high",
) -> None:
    import fitz
    from docx import Document
    from docx.shared import Pt

    dpi = quality_dpi(quality, 300)
    doc = fitz.open(src)
    out = Document()
    tmpdir = tempfile.mkdtemp(prefix="filemint-ocr-")
    try:
        all_low_conf = 0
        pages_with_text = 0
        scanned_tables = 0

        for page_index, page in enumerate(doc):
            if page_index > 0:
                out.add_page_break()
            section = out.sections[-1]
            section.page_width = Pt(page.rect.width)
            section.page_height = Pt(page.rect.height)
            section.top_margin = Pt(36)
            section.bottom_margin = Pt(36)
            section.left_margin = Pt(36)
            section.right_margin = Pt(36)

            pix = page.get_pixmap(dpi=dpi, alpha=False)
            img = os.path.join(tmpdir, f"page-{page_index + 1}.png")
            pix.save(img)
            tsv = run_tesseract_tsv(img, lang, psm="6")
            lines = parse_tsv(
                tsv, pix.width, pix.height, page.rect.width, page.rect.height
            )
            if lines:
                pages_with_text += 1
            all_low_conf += sum(
                1 for line in lines for word in line.words if 0 <= word.conf < 55
            )

            runs = table_runs(lines) if table_detection else []
            table_line_ids = {id(line) for run in runs for line in run}
            run_by_first_id = {id(run[0]): run for run in runs}
            scanned_tables += len(runs)

            prev_bottom: float | None = None
            for line in lines:
                if id(line) in run_by_first_id:
                    add_table(out, run_by_first_id[id(line)], report)
                    prev_bottom = max(
                        l.top + l.height for l in run_by_first_id[id(line)]
                    )
                    continue
                if id(line) in table_line_ids:
                    continue
                add_paragraph_from_line(out, line, prev_bottom)
                prev_bottom = line.top + line.height

            if not lines:
                p = out.add_paragraph()
                p.add_run("")

        if pages_with_text == 0:
            report["warnings"].append("OCR produced no editable text.")
        if table_detection and scanned_tables == 0:
            report["warnings"].append(
                "No scanned tables were confidently reconstructed. Table detection may need clearer grid lines or better OCR language data."
            )
        report["editableTextDetected"] = pages_with_text > 0
        report["lowConfidenceOcrAreas"] = all_low_conf
        report["notes"].append(
            "Scanned PDF rebuilt from Tesseract OCR word positions into editable DOCX text and simple tables."
        )
        out.save(dst)
    finally:
        doc.close()
        shutil.rmtree(tmpdir, ignore_errors=True)


def to_docx_image(
    src: str,
    dst: str,
    report: dict[str, Any],
    quality: str = "high",
    visual_object_format: str = "png",
) -> None:
    import fitz
    from docx import Document
    from docx.shared import Inches, Pt

    pdf = fitz.open(src)
    tmpdir = tempfile.mkdtemp(prefix="filemint-image-")
    try:
        out = Document()
        for i, page in enumerate(pdf):
            if i > 0:
                out.add_page_break()
            section = out.sections[-1]
            section.page_width = Pt(page.rect.width)
            section.page_height = Pt(page.rect.height)
            section.top_margin = Pt(0)
            section.bottom_margin = Pt(0)
            section.left_margin = Pt(0)
            section.right_margin = Pt(0)
            ext = "jpg" if visual_object_format == "jpg" else "png"
            pix = page.get_pixmap(dpi=quality_dpi(quality, 180), alpha=False)
            img = os.path.join(tmpdir, f"page-{i + 1}.{ext}")
            pix.save(img)
            out.add_picture(img, width=Inches(page.rect.width / 72.0))
        out.save(dst)
        report["editableTextDetected"] = False
        report["nonEditableVisualFallback"] = True
        report["warnings"].append(
            "Image-only mode creates non-editable page pictures. Use it only when editable reconstruction fails."
        )
    finally:
        pdf.close()
        shutil.rmtree(tmpdir, ignore_errors=True)
