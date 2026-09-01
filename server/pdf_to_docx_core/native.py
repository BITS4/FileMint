"""Native PDF text extraction and editable DOCX reconstruction."""

from __future__ import annotations

from dataclasses import dataclass
from statistics import median
from typing import Any


@dataclass
class PdfSpan:
    text: str
    font: str
    size: float
    flags: int
    color: int
    bbox: tuple[float, float, float, float]


@dataclass
class PdfLine:
    spans: list[PdfSpan]
    left: float
    top: float
    right: float
    bottom: float


def span_text(spans: list[PdfSpan]) -> str:
    return "".join(span.text for span in spans)


def collect_pdf_lines(page: Any) -> list[PdfLine]:
    raw = page.get_text("dict", sort=True)
    lines: list[PdfLine] = []
    for block in raw.get("blocks", []):
        if block.get("type") != 0:
            continue
        for line in block.get("lines", []):
            spans: list[PdfSpan] = []
            for span in line.get("spans", []):
                text = span.get("text", "")
                if not text:
                    continue
                bbox = tuple(
                    float(x) for x in span.get("bbox", line.get("bbox", (0, 0, 0, 0)))
                )
                spans.append(
                    PdfSpan(
                        text=text,
                        font=str(span.get("font", "")),
                        size=float(span.get("size", 11.0) or 11.0),
                        flags=int(span.get("flags", 0) or 0),
                        color=int(span.get("color", 0) or 0),
                        bbox=bbox,  # type: ignore[arg-type]
                    )
                )
            if not spans or not span_text(spans).strip():
                continue
            left, top, right, bottom = (
                float(x) for x in line.get("bbox", (0, 0, 0, 0))
            )
            lines.append(
                PdfLine(spans=spans, left=left, top=top, right=right, bottom=bottom)
            )
    return sorted(lines, key=lambda line: (line.top, line.left))


def vertical_overlap(a: PdfLine, b: PdfLine) -> float:
    overlap = min(a.bottom, b.bottom) - max(a.top, b.top)
    if overlap <= 0:
        return 0.0
    return overlap / max(1.0, min(a.bottom - a.top, b.bottom - b.top))


def merge_visual_rows(lines: list[PdfLine]) -> list[PdfLine]:
    rows: list[list[PdfLine]] = []
    for line in lines:
        placed = False
        center = (line.top + line.bottom) / 2
        height = max(1.0, line.bottom - line.top)
        for row in rows:
            row_top = min(item.top for item in row)
            row_bottom = max(item.bottom for item in row)
            row_center = (row_top + row_bottom) / 2
            if abs(center - row_center) <= max(2.5, height * 0.42) or any(
                vertical_overlap(line, item) > 0.68 for item in row
            ):
                row.append(line)
                placed = True
                break
        if not placed:
            rows.append([line])

    merged: list[PdfLine] = []
    for row in rows:
        ordered = sorted(row, key=lambda item: item.left)
        spans: list[PdfSpan] = []
        for item in ordered:
            spans.extend(item.spans)
        merged.append(
            PdfLine(
                spans=spans,
                left=min(item.left for item in ordered),
                top=min(item.top for item in ordered),
                right=max(item.right for item in ordered),
                bottom=max(item.bottom for item in ordered),
            )
        )
    return sorted(merged, key=lambda line: (line.top, line.left))


def word_font_name(pdf_font: str) -> str:
    font = pdf_font.upper()
    if "TT" in font or "MONO" in font or "COURIER" in font:
        return "Courier New"
    return "Times New Roman"


def span_is_bold(span: PdfSpan) -> bool:
    font = span.font.upper()
    return bool(span.flags & 16) or "BOLD" in font or "CMBX" in font or "BX" in font


def span_is_italic(span: PdfSpan) -> bool:
    font = span.font.upper()
    return (
        bool(span.flags & 2)
        or "ITAL" in font
        or "CMMI" in font
        or "CMTI" in font
        or "MI" in font
    )


def set_run_color(run: Any, color: int) -> None:
    if color == 0:
        return
    try:
        from docx.shared import RGBColor

        run.font.color.rgb = RGBColor(
            (color >> 16) & 255, (color >> 8) & 255, color & 255
        )
    except Exception:
        pass


def xml_compatible_text(text: str) -> str:
    return "".join(ch for ch in text if ch in "\t\n\r" or ord(ch) >= 32)


def append_pdf_span_run(
    paragraph: Any, span: PdfSpan, row: PdfLine, dominant_size: float
) -> None:
    from docx.shared import Pt

    text = xml_compatible_text(span.text)
    if not text:
        return
    run = paragraph.add_run(text)
    run.font.name = word_font_name(span.font)
    run.font.size = Pt(max(5.0, min(28.0, span.size)))
    run.bold = span_is_bold(span)
    run.italic = span_is_italic(span)
    set_run_color(run, span.color)

    if span.size < dominant_size * 0.82:
        span_mid = (span.bbox[1] + span.bbox[3]) / 2
        row_mid = (row.top + row.bottom) / 2
        if span_mid < row_mid - dominant_size * 0.08:
            run.font.superscript = True
        elif span_mid > row_mid + dominant_size * 0.08:
            run.font.subscript = True


def estimated_gap_spaces(gap_pt: float, size_pt: float) -> int:
    if gap_pt <= max(1.0, size_pt * 0.12):
        return 0
    return max(1, min(24, round(gap_pt / max(2.5, size_pt * 0.33))))


def append_pdf_line(
    doc: Any, line: PdfLine, prev_bottom: float | None, left_margin_pt: float
) -> int:
    from docx.shared import Pt

    paragraph = doc.add_paragraph()
    fmt = paragraph.paragraph_format
    fmt.left_indent = Pt(max(0.0, line.left - left_margin_pt))
    fmt.space_after = Pt(0)
    fmt.line_spacing = 1.0
    if prev_bottom is None:
        fmt.space_before = Pt(0)
    else:
        gap = max(0.0, line.top - prev_bottom)
        fmt.space_before = Pt(min(40.0, gap))

    sizes = [span.size for span in line.spans if span.text.strip()]
    dominant_size = median(sizes) if sizes else 11.0
    last_right: float | None = None
    emitted_chars = 0
    for span in sorted(line.spans, key=lambda s: (s.bbox[0], s.bbox[1])):
        text = span.text
        if last_right is not None:
            gap = span.bbox[0] - last_right
            spaces = estimated_gap_spaces(gap, dominant_size)
            if spaces and not text.startswith(" "):
                paragraph.add_run(" " * spaces)
                emitted_chars += spaces
        append_pdf_span_run(paragraph, span, line, dominant_size)
        emitted_chars += len(text)
        last_right = max(last_right or span.bbox[2], span.bbox[2])
    return emitted_chars


def to_docx_digital_text_flow(src: str, dst: str, report: dict[str, Any]) -> None:
    import fitz
    from docx import Document
    from docx.shared import Pt

    pdf = fitz.open(src)
    out = Document()
    try:
        total_lines = 0
        total_chars = 0
        for page_index, page in enumerate(pdf):
            if page_index == 0:
                section = out.sections[-1]
            else:
                out.add_page_break()
                section = out.sections[-1]
            section.page_width = Pt(page.rect.width)
            section.page_height = Pt(page.rect.height)
            section.left_margin = Pt(72)
            section.right_margin = Pt(72)
            section.top_margin = Pt(72)
            section.bottom_margin = Pt(72)

            rows = merge_visual_rows(collect_pdf_lines(page))
            prev_bottom: float | None = None
            for row in rows:
                total_chars += append_pdf_line(out, row, prev_bottom, 72.0)
                total_lines += 1
                prev_bottom = row.bottom

        report["resolvedMode"] = "premium-digital-flow"
        report["editableTextDetected"] = True
        report["editableTextBoxes"] = total_lines
        report["editableCharacters"] = total_chars
        report["ocrTextCandidates"] = total_lines
        report["textCoverageEstimate"] = 100 if total_lines else 0
        report["notes"].append(
            "Premium digital text-flow rebuilt the native PDF text layer as editable Word text with inferred word spacing, line gaps and page breaks."
        )
        out.save(dst)
    finally:
        pdf.close()
