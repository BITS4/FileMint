"""Scanned transcript table reconstruction for Word output."""

from __future__ import annotations

import os
from typing import Any

from .docx import VML_NS, line_text_signal
from .models import LineBox
from .transcript import (
    detect_transcript_grid_geometry,
    join_positioned_words,
    normalize_grade_cell,
    normalize_transcript_course_cell,
    set_cell_margins,
    set_cell_width,
    set_row_height,
    set_table_cell_text,
    set_table_layout,
    transcript_row_centers,
    words_from_lines,
    words_in_box,
)


def build_scanned_table_page(
    doc: Any,
    scan_png: str,
    lines: list[LineBox],
    page_width_pt: float,
    page_height_pt: float,
    report: dict[str, Any],
) -> None:
    from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import parse_xml
    from docx.shared import Pt
    from PIL import Image

    section = doc.sections[-1]
    section.page_width = Pt(page_width_pt)
    section.page_height = Pt(page_height_pt)
    section.top_margin = Pt(28)
    section.bottom_margin = Pt(26)
    section.left_margin = Pt(34)
    section.right_margin = Pt(34)

    img = Image.open(scan_png).convert("RGB")
    width_px, height_px = img.size
    words = words_from_lines(lines)
    grid_geometry = detect_transcript_grid_geometry(img, page_width_pt, page_height_pt)
    left_margin_pt = 34.0

    header_lines = [
        line
        for line in lines
        if line.top < height_px * 0.17 and line.conf >= 35 and line_text_signal(line)
    ]
    columns = [[], [], []]
    for line in header_lines:
        center = line.left + line.width / 2
        idx = 0 if center < width_px * 0.35 else 1 if center < width_px * 0.66 else 2
        columns[idx].append(line)
    header = doc.add_table(rows=1, cols=3)
    header.alignment = WD_TABLE_ALIGNMENT.CENTER
    header.autofit = False
    for idx, cell in enumerate(header.rows[0].cells):
        text = "\n".join(
            line.text for line in sorted(columns[idx], key=lambda l: l.top)[:5]
        )
        set_cell_width(cell, 165)
        set_table_cell_text(cell, text, size=7.2, bold=idx < 3, align="center")

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_before = Pt(4)
    title.paragraph_format.space_after = Pt(3)
    run = title.add_run("Student Personal Information")
    run.font.name = "Arial"
    run.font.size = Pt(11)
    run.bold = True

    info = doc.add_table(rows=2, cols=2)
    info.style = "Table Grid"
    info.alignment = WD_TABLE_ALIGNMENT.CENTER
    info.autofit = False
    info_col_w = [250, 250]
    for row in info.rows:
        set_row_height(row, 20)
        for i, cell in enumerate(row.cells):
            set_cell_width(cell, info_col_w[i])
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

    def box_text(left_r: float, top_r: float, right_r: float, bottom_r: float) -> str:
        return join_positioned_words(
            words_in_box(
                words,
                width_px * left_r,
                height_px * top_r,
                width_px * right_r,
                height_px * bottom_r,
            )
        )

    set_table_cell_text(info.cell(0, 0), box_text(0.06, 0.205, 0.47, 0.245), size=7.8)
    set_table_cell_text(info.cell(0, 1), box_text(0.47, 0.205, 0.93, 0.245), size=7.8)
    set_table_cell_text(info.cell(1, 0), box_text(0.06, 0.235, 0.47, 0.275), size=7.8)
    set_table_cell_text(info.cell(1, 1), box_text(0.47, 0.235, 0.93, 0.275), size=7.8)

    academic = doc.add_paragraph()
    academic.alignment = WD_ALIGN_PARAGRAPH.CENTER
    academic.paragraph_format.space_before = Pt(6)
    academic.paragraph_format.space_after = Pt(
        grid_geometry["preTableSpacerPt"] if grid_geometry else 4
    )
    arun = academic.add_run("Academic Record 2024 - 2026")
    arun.font.name = "Arial"
    arun.font.size = Pt(9.5)
    arun.bold = True

    if grid_geometry and len(grid_geometry["yPositionsPx"]) >= 25:
        y_positions_px = grid_geometry["yPositionsPx"]
        row_bounds = [
            (y_positions_px[i], y_positions_px[i + 1])
            for i in range(1, min(24, len(y_positions_px) - 1))
        ]
    else:
        row_centers = transcript_row_centers(words, width_px, height_px)
        if not row_centers:
            row_centers = [
                height_px * 0.365 + i * height_px * 0.0214 for i in range(23)
            ]
        row_bounds = []
        header_top = max(height_px * 0.30, row_centers[0] - 130)
        first_boundary = (header_top + row_centers[0]) / 2
        for i, center in enumerate(row_centers):
            top = first_boundary if i == 0 else (row_centers[i - 1] + center) / 2
            bottom = (
                height_px * 0.86
                if i == len(row_centers) - 1
                else (center + row_centers[i + 1]) / 2
            )
            row_bounds.append((top, bottom))

    table = doc.add_table(rows=len(row_bounds) + 1, cols=5)
    table.style = "Table Grid"
    table.alignment = (
        WD_TABLE_ALIGNMENT.LEFT if grid_geometry else WD_TABLE_ALIGNMENT.CENTER
    )
    table.autofit = False
    col_widths_pt = (
        grid_geometry["colWidthsPt"][:5]
        if grid_geometry and len(grid_geometry["colWidthsPt"]) >= 5
        else [34, 205, 86, 86, 90]
    )
    if grid_geometry:
        set_table_layout(
            table,
            width_pt=grid_geometry["tableWidthPt"],
            indent_pt=grid_geometry["tableLeftPt"] - left_margin_pt,
        )
        report["notes"].append(
            "Transcript table geometry was measured from the source scan and applied to Word column widths, row heights and table position."
        )
    else:
        set_table_layout(table, width_pt=sum(col_widths_pt))
    headers = ["#", "Course Title", "Grade 9", "Grade 10", "Grade 11\n(half-year)"]
    for i, cell in enumerate(table.rows[0].cells):
        set_cell_width(cell, col_widths_pt[i])
        set_cell_margins(cell, 1.4 if grid_geometry else 2.0)
        set_table_cell_text(
            cell,
            headers[i],
            size=7.8 if grid_geometry else 8.5,
            bold=True,
            align="center" if i != 1 else "left",
        )
    header_height = (
        grid_geometry["rowHeightsPt"][0]
        if grid_geometry and grid_geometry["rowHeightsPt"]
        else 26
    )
    set_row_height(
        table.rows[0], header_height, "exact" if grid_geometry else "atLeast"
    )

    if grid_geometry and len(grid_geometry["xPositionsPx"]) >= 6:
        x_positions_px = grid_geometry["xPositionsPx"][:6]
        col_bounds = [(x_positions_px[i], x_positions_px[i + 1]) for i in range(5)]
    else:
        col_bounds = [
            (width_px * 0.20, width_px * 0.25),
            (width_px * 0.25, width_px * 0.535),
            (width_px * 0.535, width_px * 0.660),
            (width_px * 0.660, width_px * 0.775),
            (width_px * 0.775, width_px * 0.890),
        ]
    for r_idx, (top, bottom) in enumerate(row_bounds, start=1):
        row = table.rows[r_idx]
        if grid_geometry and len(grid_geometry["rowHeightsPt"]) > r_idx:
            set_row_height(row, grid_geometry["rowHeightsPt"][r_idx], "exact")
        else:
            set_row_height(row, 18)
        for c_idx, cell in enumerate(row.cells):
            set_cell_width(cell, col_widths_pt[c_idx])
            set_cell_margins(cell, 1.2 if grid_geometry else 2.0)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            if c_idx == 0:
                text = str(r_idx)
            else:
                left, right = col_bounds[c_idx]
                text = join_positioned_words(
                    words_in_box(words, left, top, right, bottom)
                )
                if c_idx == 1:
                    text = normalize_transcript_course_cell(text, r_idx)
                if c_idx >= 2:
                    text = normalize_grade_cell(text)
            set_table_cell_text(
                cell,
                text,
                size=(
                    7.0
                    if grid_geometry and c_idx == 1
                    else 6.8 if grid_geometry else 7.8 if c_idx == 1 else 7.6
                ),
                bold=False,
                align="left" if c_idx == 1 else "center",
            )

    footer = doc.add_paragraph()
    footer.paragraph_format.space_before = Pt(5)
    footer.paragraph_format.space_after = Pt(0)
    footer.alignment = WD_ALIGN_PARAGRAPH.LEFT
    frun = footer.add_run(
        'Principal of Tajik - Russian Lyceum - Boarding "Hotamov Abdufattоh"'
    )
    frun.font.name = "Arial"
    frun.font.size = Pt(7.5)

    tmpdir = os.path.dirname(scan_png)
    seal_path = os.path.join(tmpdir, "seal-signature.png")
    crop = img.crop(
        (
            int(width_px * 0.30),
            int(height_px * 0.84),
            int(width_px * 0.80),
            min(height_px, int(height_px * 0.985)),
        )
    ).convert("RGBA")
    pixels = crop.load()
    for y in range(crop.height):
        for x in range(crop.width):
            r, g, b, _a = pixels[x, y]
            colorful = max(r, g, b) - min(r, g, b) > 28
            blue_mark = b > r + 10 and b > g - 8 and b < 245
            if not (colorful and blue_mark):
                pixels[x, y] = (255, 255, 255, 0)
    crop.save(seal_path)
    r_id, _image = doc.part.get_or_add_image(seal_path)
    seal_left_pt = page_width_pt * 0.30
    seal_top_pt = page_height_pt * 0.825
    seal_width_pt = page_width_pt * 0.50
    seal_height_pt = page_height_pt * 0.155
    footer._p.append(
        parse_xml(
            f"""
            <w:r {VML_NS}>
              <w:pict>
                <v:shape id="FileMintTranscriptSeal" type="#_x0000_t75"
                  style="position:absolute;margin-left:{seal_left_pt:.2f}pt;margin-top:{seal_top_pt:.2f}pt;width:{seal_width_pt:.2f}pt;height:{seal_height_pt:.2f}pt;z-index:2500;mso-position-horizontal:absolute;mso-position-horizontal-relative:page;mso-position-vertical:absolute;mso-position-vertical-relative:page"
                  stroked="f" filled="f" o:allowincell="f">
                  <v:imagedata r:id="{r_id}" o:title="FileMint seal and signature"/>
                </v:shape>
              </w:pict>
            </w:r>
            """
        )
    )

    report["notes"].append(
        "Premium editable table mode rebuilt the scanned transcript as a real Word table; stamp and signature regions were kept as images."
    )
