"""Positioned and exact-visual DOCX page rendering."""

from __future__ import annotations

import math
from statistics import median
from typing import Any
from xml.sax.saxutils import escape

from .docx import (
    VML_NS,
    contains_rtl,
    ocr_font_attrs,
    ocr_font_size,
    set_paragraph_bidi,
    set_run_font,
)
from .models import LineBox, VisualFragment, VisualRule
from .native import (
    PdfLine,
    collect_pdf_lines,
    estimated_gap_spaces,
    merge_visual_rows,
    xml_compatible_text,
)
from .selection import line_should_be_bold


def append_exact_visual_page(
    doc: Any,
    page_index: int,
    visual_png: str,
    editable_lines: list[LineBox],
    page_width_pt: float,
    page_height_pt: float,
    hidden_text: bool = False,
) -> None:
    from docx.oxml import parse_xml
    from docx.shared import Pt

    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1

    # A purely floating VML page image does not reserve vertical layout space.
    # LibreOffice/Word can then anchor multiple pages to one physical page,
    # especially after page breaks. The inline page image below occupies the
    # page, while the VML text boxes remain editable overlays on that page.
    paragraph.add_run().add_picture(
        visual_png,
        width=Pt(max(1.0, page_width_pt - 0.8)),
        height=Pt(max(1.0, page_height_pt - 0.8)),
    )

    scale_x = page_width_pt / max(
        1.0, editable_lines[0].page_width_px if editable_lines else page_width_pt
    )
    scale_y = page_height_pt / max(
        1.0, editable_lines[0].page_height_px if editable_lines else page_height_pt
    )
    for idx, line in enumerate(editable_lines):
        x = max(0.0, line.left * scale_x - 1.0)
        y = max(0.0, line.top * scale_y - 1.0)
        w = min(
            max(14.0, line.width * scale_x * 1.12 + 10.0),
            max(14.0, page_width_pt - x - 2.0),
        )
        h = max(8.0, line.height * scale_y * 1.9)
        font_size = ocr_font_size(line)
        text = escape(line.text)
        align = "right" if contains_rtl(line.text) else "left"
        bidi = "<w:bidi/>" if contains_rtl(line.text) else ""
        bold = "<w:b/>" if line_should_be_bold(line) else ""
        vanish = "<w:vanish/>" if hidden_text else ""
        color = '<w:color w:val="FFFFFF"/>' if hidden_text else ""
        ascii_font, east_asia_font, complex_font = ocr_font_attrs(line.text)
        paragraph._p.append(
            parse_xml(
                f"""
            <w:r {VML_NS}>
              <w:pict>
                <v:shape id="FileMintText{page_index}_{idx}" type="#_x0000_t202"
                  style="position:absolute;margin-left:{x:.2f}pt;margin-top:{y:.2f}pt;width:{w:.2f}pt;height:{h:.2f}pt;z-index:{idx + 1};mso-position-horizontal:absolute;mso-position-horizontal-relative:page;mso-position-vertical:absolute;mso-position-vertical-relative:page"
                  stroked="f" filled="f" o:allowincell="f">
                  <v:textbox inset="0,0,0,0">
                    <w:txbxContent>
                      <w:p>
                        <w:pPr><w:jc w:val="{align}"/>{bidi}</w:pPr>
                        <w:r>
                          <w:rPr>
                            <w:rFonts w:ascii="{ascii_font}" w:hAnsi="{ascii_font}" w:eastAsia="{east_asia_font}" w:cs="{complex_font}"/>
                            {bold}
                            {vanish}
                            {color}
                            <w:sz w:val="{int(font_size * 2)}"/>
                            <w:szCs w:val="{int(font_size * 2)}"/>
                          </w:rPr>
                          <w:t xml:space="preserve">{text}</w:t>
                        </w:r>
                      </w:p>
                    </w:txbxContent>
                  </v:textbox>
                </v:shape>
              </w:pict>
            </w:r>
            """
            )
        )


def append_positioned_page(
    doc: Any,
    page_index: int,
    editable_lines: list[LineBox],
    page_width_pt: float,
    page_height_pt: float,
    fragments: list[VisualFragment],
    rules: list[VisualRule],
) -> None:
    from docx.oxml import parse_xml
    from docx.shared import Pt, RGBColor

    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = Pt(max(1.0, page_height_pt - 1.0))
    spacer = paragraph.add_run(" ")
    spacer.font.size = Pt(1)
    try:
        spacer.font.color.rgb = RGBColor(255, 255, 255)
    except Exception:
        pass

    for idx, rule in enumerate(rules):
        x = max(0.0, rule.left)
        y = max(0.0, rule.top)
        w = min(max(0.35, rule.width), max(0.35, page_width_pt - x))
        h = min(max(0.35, rule.height), max(0.35, page_height_pt - y))
        paragraph._p.append(
            parse_xml(
                f"""
                <w:r {VML_NS}>
                  <w:pict>
                    <v:rect id="FileMintRule{page_index}_{idx}"
                      style="position:absolute;margin-left:{x:.2f}pt;margin-top:{y:.2f}pt;width:{w:.2f}pt;height:{h:.2f}pt;z-index:{10 + idx};mso-position-horizontal:absolute;mso-position-horizontal-relative:page;mso-position-vertical:absolute;mso-position-vertical-relative:page"
                      fillcolor="{rule.color}" stroked="f" filled="t" o:allowincell="f"/>
                  </w:pict>
                </w:r>
                """
            )
        )

    for idx, fragment in enumerate(fragments):
        if fragment.width <= 0 or fragment.height <= 0:
            continue
        r_id, _image = doc.part.get_or_add_image(fragment.path)
        x = max(0.0, fragment.left)
        y = max(0.0, fragment.top)
        w = min(max(0.5, fragment.width), max(0.5, page_width_pt - x))
        h = min(max(0.5, fragment.height), max(0.5, page_height_pt - y))
        paragraph._p.append(
            parse_xml(
                f"""
                <w:r {VML_NS}>
                  <w:pict>
                    <v:shape id="FileMintFragment{page_index}_{idx}" type="#_x0000_t75"
                      style="position:absolute;margin-left:{x:.2f}pt;margin-top:{y:.2f}pt;width:{w:.2f}pt;height:{h:.2f}pt;z-index:{100 + idx};mso-position-horizontal:absolute;mso-position-horizontal-relative:page;mso-position-vertical:absolute;mso-position-vertical-relative:page"
                      stroked="f" filled="f" o:allowincell="f">
                      <v:imagedata r:id="{r_id}" o:title="FileMint visual fragment"/>
                    </v:shape>
                  </w:pict>
                </w:r>
                """
            )
        )

    scale_x = page_width_pt / max(
        1.0, editable_lines[0].page_width_px if editable_lines else page_width_pt
    )
    scale_y = page_height_pt / max(
        1.0, editable_lines[0].page_height_px if editable_lines else page_height_pt
    )
    for idx, line in enumerate(editable_lines):
        x = max(0.0, line.left * scale_x - 1.0)
        y = max(0.0, line.top * scale_y - 1.0)
        w = min(
            max(14.0, line.width * scale_x * 1.12 + 10.0),
            max(14.0, page_width_pt - x - 2.0),
        )
        h = max(8.0, line.height * scale_y * 1.9)
        font_size = ocr_font_size(line)
        text = escape(line.text)
        align = "right" if contains_rtl(line.text) else "left"
        bidi = "<w:bidi/>" if contains_rtl(line.text) else ""
        bold = "<w:b/>" if line_should_be_bold(line) else ""
        ascii_font, east_asia_font, complex_font = ocr_font_attrs(line.text)
        paragraph._p.append(
            parse_xml(
                f"""
                <w:r {VML_NS}>
                  <w:pict>
                    <v:shape id="FileMintText{page_index}_{idx}" type="#_x0000_t202"
                      style="position:absolute;margin-left:{x:.2f}pt;margin-top:{y:.2f}pt;width:{w:.2f}pt;height:{h:.2f}pt;z-index:{1000 + idx};mso-position-horizontal:absolute;mso-position-horizontal-relative:page;mso-position-vertical:absolute;mso-position-vertical-relative:page"
                      stroked="f" filled="f" o:allowincell="f">
                      <v:textbox inset="0,0,0,0">
                        <w:txbxContent>
                          <w:p>
                            <w:pPr><w:jc w:val="{align}"/>{bidi}</w:pPr>
                            <w:r>
                              <w:rPr>
                                <w:rFonts w:ascii="{ascii_font}" w:hAnsi="{ascii_font}" w:eastAsia="{east_asia_font}" w:cs="{complex_font}"/>
                                {bold}
                                <w:sz w:val="{int(font_size * 2)}"/>
                                <w:szCs w:val="{int(font_size * 2)}"/>
                              </w:rPr>
                              <w:t xml:space="preserve">{text}</w:t>
                            </w:r>
                          </w:p>
                        </w:txbxContent>
                      </v:textbox>
                    </v:shape>
                  </w:pict>
                </w:r>
                """
            )
        )


def pdf_line_text_with_gaps(line: PdfLine) -> str:
    spans = sorted(line.spans, key=lambda s: (s.bbox[0], s.bbox[1]))
    sizes = [span.size for span in spans if span.text.strip()]
    dominant_size = median(sizes) if sizes else 11.0
    parts: list[str] = []
    last_right: float | None = None
    for span in spans:
        text = xml_compatible_text(span.text)
        if not text:
            continue
        if last_right is not None:
            gap = span.bbox[0] - last_right
            spaces = estimated_gap_spaces(gap, dominant_size)
            if (
                spaces
                and parts
                and not parts[-1].endswith(" ")
                and not text.startswith(" ")
            ):
                parts.append(" " * spaces)
        parts.append(text)
        last_right = max(last_right or span.bbox[2], span.bbox[2])
    return "".join(parts).strip()


def native_pdf_line_boxes(
    page: Any, page_width_px: float, page_height_px: float
) -> list[LineBox]:
    sx = page_width_px / max(1.0, page.rect.width)
    sy = page_height_px / max(1.0, page.rect.height)
    out: list[LineBox] = []
    for line in merge_visual_rows(collect_pdf_lines(page)):
        text = pdf_line_text_with_gaps(line)
        if not text or not any(ch.isalnum() for ch in text):
            continue
        left = line.left * sx
        top = line.top * sy
        width = max(1.0, (line.right - line.left) * sx)
        height = max(1.0, (line.bottom - line.top) * sy)
        out.append(
            LineBox(
                text=text,
                words=[],
                left=left,
                top=top,
                width=width,
                height=height,
                conf=99.0,
                page_width_px=page_width_px,
                page_height_px=page_height_px,
                page_width_pt=page.rect.width,
                page_height_pt=page.rect.height,
                segments=[(left, left + width, text)],
            )
        )
    return sorted(out, key=lambda line: (line.top, line.left))


def image_backed_text_layer_likely(info: dict[str, Any]) -> bool:
    pages = int(info.get("pages", 0) or 0)
    if pages <= 0:
        return False
    details = info.get("pageDetails", []) or []
    backed = [
        page
        for page in details
        if float(page.get("maxImageCoverage", 0) or 0) >= 0.65
        and int(page.get("textCharacters", 0) or 0) >= 25
    ]
    return len(backed) >= max(1, math.ceil(pages * 0.5))


def image_backed_text_layer_needs_ocr(info: dict[str, Any]) -> bool:
    for page in info.get("pageDetails", []) or []:
        if (
            float(page.get("maxImageCoverage", 0) or 0) >= 0.65
            and int(page.get("textCharacters", 0) or 0) < 200
        ):
            return True
    return False


def append_linebox_flow_paragraph(
    doc: Any,
    line: LineBox,
    prev_bottom: float | None,
    left_margin_pt: float,
    top_margin_pt: float,
) -> int:
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt

    text = xml_compatible_text(" ".join(line.text.split()))
    if not text:
        return 0

    scale_x = line.page_width_pt / max(1.0, line.page_width_px)
    scale_y = line.page_height_pt / max(1.0, line.page_height_px)
    left_pt = line.left * scale_x
    right_pt = (line.left + line.width) * scale_x
    top_pt = line.top * scale_y
    bottom_pt = (line.top + line.height) * scale_y
    page_center = line.page_width_pt / 2
    line_center = (left_pt + right_pt) / 2

    paragraph = doc.add_paragraph()
    fmt = paragraph.paragraph_format
    fmt.left_indent = Pt(max(0.0, left_pt - left_margin_pt))
    fmt.right_indent = Pt(0)
    fmt.space_after = Pt(0)
    if prev_bottom is None:
        fmt.space_before = Pt(max(0.0, min(42.0, top_pt - top_margin_pt)))
    else:
        gap = max(0.0, top_pt - prev_bottom)
        fmt.space_before = Pt(min(28.0, gap))
    fmt.line_spacing = 1.0

    if contains_rtl(text):
        set_paragraph_bidi(paragraph)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    elif (
        abs(line_center - page_center) <= 35
        and (right_pt - left_pt) <= line.page_width_pt * 0.76
    ):
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        fmt.left_indent = Pt(0)
    else:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT

    run = paragraph.add_run(text)
    set_run_font(run, "Times New Roman")
    font_size = max(7.0, min(16.0, (bottom_pt - top_pt) * 0.92))
    run.font.size = Pt(font_size)
    run.bold = line_should_be_bold(line) or font_size >= 12.5
    return len(text)
