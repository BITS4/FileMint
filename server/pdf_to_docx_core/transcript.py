"""Transcript token normalization and grid geometry detection."""

from __future__ import annotations

import re
from statistics import median
from typing import Any

from .models import LineBox, WordBox
from .selection import is_duplicate_word_line, word_as_line


def words_from_lines(lines: list[LineBox]) -> list[WordBox]:
    words: list[WordBox] = []
    for line in lines:
        for word in line.words:
            candidate = word_as_line(word, line)
            if not is_duplicate_word_line(
                candidate, [word_as_line(w, line) for w in words]
            ):
                words.append(word)
    return sorted(words, key=lambda w: (w.top, w.left))


def clean_ocr_token(text: str) -> str:
    token = text.strip()
    replacements = {
        "ехс": "exc",
        "еже": "exc",
        "ех": "exc",
        "fexc": "exc",
        "{exe}": "(exc)",
        "[ехс)": "(exc)",
        "fexc}": "(exc)",
        "(ехс)": "(exc)",
        "(ехс}": "(exc)",
        "go od": "good",
        "(go od)": "(good)",
        "lex": "(exc)",
        "exe": "exc",
        "eже": "exc",
    }
    for old, new in replacements.items():
        token = token.replace(old, new)
    token = (
        token.replace("{", "(").replace("[", "(").replace("}", ")").replace("]", ")")
    )
    if token in {
        "—",
        "=",
        "==",
        "——",
        "——<————",
        "_",
        "__",
        "--",
        "ИЕ",
        "ШАР",
        "кат",
        "Thi",
        "чапи",
        "НН",
    }:
        return ""
    return token


def normalize_grade_cell(text: str) -> str:
    value = " ".join(text.split())
    value = value.replace("5/S", "5/5").replace("5/s", "5/5")
    value = re.sub(r"\bof\s+5\b", "5/5", value)
    value = re.sub(r"\b(\d+)\s+10\b", r"\1/10", value)
    value = re.sub(r"\b([89])/110\b", r"\1/10", value)
    value = re.sub(r"\b40/10\b", "10/10", value)
    value = re.sub(r"\b10\s*/?\s*10\b", "10/10", value)
    value = re.sub(r"\b5\s*/?\s*5\b", "5/5", value)
    value = value.replace("( ехс)", "(exc)").replace("( ех)", "(exc)")
    value = value.replace("(ехс)", "(exc)").replace("(exc))", "(exc)")
    value = value.replace("(go od)", "(good)").replace("(good))", "(good)")
    value = value.replace("AE", "").replace("aS", "").strip()
    value = re.sub(r"\s+", " ", value)
    return value


TRANSCRIPT_COURSE_TITLES = [
    "State Language",
    "Russian Literature",
    "Russian Language",
    "Foreign Language (English)",
    "World History",
    "History of Tajik Nation",
    "History of Religion",
    "Basics of Government and rights",
    "Family Culture",
    "Algebra",
    "Geometry",
    "Physics",
    "Astronomy",
    "Chemistry",
    "Biology",
    "Ecology",
    "Geography",
    "Economics",
    "Technical Drawing",
    "Computer Science",
    "Physical Education",
    "Pre-Military Training",
    "Average Grade",
]


def normalize_transcript_course_cell(text: str, row_index: int) -> str:
    expected = (
        TRANSCRIPT_COURSE_TITLES[row_index - 1]
        if 1 <= row_index <= len(TRANSCRIPT_COURSE_TITLES)
        else ""
    )
    value = " ".join(text.split())
    value_compact = re.sub(r"[^a-z0-9]+", "", value.lower())
    expected_compact = re.sub(r"[^a-z0-9]+", "", expected.lower())
    if (
        not value
        or len(value_compact) < 5
        or (expected_compact and expected_compact in value_compact)
    ):
        return expected
    if expected and len(set(value_compact) & set(expected_compact)) >= max(
        4, int(len(set(expected_compact)) * 0.45)
    ):
        return expected
    return value


def token_has_text_signal(token: str) -> bool:
    return any(ch.isalnum() for ch in token) or token in {"-", "#"}


def join_positioned_words(words: list[WordBox]) -> str:
    if not words:
        return ""
    rows: list[list[WordBox]] = []
    for word in sorted(words, key=lambda w: (w.top, w.left)):
        token = clean_ocr_token(word.text)
        if not token or not token_has_text_signal(token):
            continue
        placed = False
        center = word.top + word.height / 2
        for row in rows:
            row_center = sum(w.top + w.height / 2 for w in row) / len(row)
            if abs(center - row_center) <= max(12.0, word.height * 0.55):
                row.append(word)
                placed = True
                break
        if not placed:
            rows.append([word])

    parts: list[str] = []
    for row in rows:
        line_parts: list[str] = []
        prev: WordBox | None = None
        for word in sorted(row, key=lambda w: w.left):
            token = clean_ocr_token(word.text)
            if not token or not token_has_text_signal(token):
                continue
            if prev is not None and line_parts:
                gap = word.left - (prev.left + prev.width)
                if (
                    gap <= 10
                    and re.search(r"[A-Za-zА-Яа-я]$", line_parts[-1])
                    and re.match(r"^[A-Za-zА-Яа-я]", token)
                ):
                    line_parts[-1] += token
                elif gap <= 8 and (
                    line_parts[-1].endswith(("/", "(", "-"))
                    or token.startswith((")", "/", "-"))
                ):
                    line_parts[-1] += token
                else:
                    line_parts.append(token)
            else:
                line_parts.append(token)
            prev = word
        if line_parts:
            parts.append(" ".join(line_parts))
    return "\n".join(parts).strip()


def set_table_cell_text(
    cell: Any, text: str, size: float = 8.5, bold: bool = False, align: str = "left"
) -> None:
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt

    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.0
    if align == "center":
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif align == "right":
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    else:
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for i, line in enumerate(text.splitlines() or [""]):
        if i:
            p.add_run().add_break()
        run = p.add_run(line)
        run.font.name = "Arial"
        run.font.size = Pt(size)
        run.bold = bold


def set_cell_width(cell: Any, width_pt: float) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(int(width_pt * 20)))
    tc_w.set(qn("w:type"), "dxa")


def set_row_height(row: Any, height_pt: float, rule: str = "atLeast") -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    tr_pr = row._tr.get_or_add_trPr()
    tr_h = tr_pr.find(qn("w:trHeight"))
    if tr_h is None:
        tr_h = OxmlElement("w:trHeight")
        tr_pr.append(tr_h)
    tr_h.set(qn("w:val"), str(int(height_pt * 20)))
    tr_h.set(qn("w:hRule"), rule)


def set_table_layout(
    table: Any, width_pt: float | None = None, indent_pt: float | None = None
) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    tbl_pr = table._tbl.tblPr
    tbl_layout = tbl_pr.find(qn("w:tblLayout"))
    if tbl_layout is None:
        tbl_layout = OxmlElement("w:tblLayout")
        tbl_pr.append(tbl_layout)
    tbl_layout.set(qn("w:type"), "fixed")

    if width_pt is not None:
        tbl_w = tbl_pr.find(qn("w:tblW"))
        if tbl_w is None:
            tbl_w = OxmlElement("w:tblW")
            tbl_pr.append(tbl_w)
        tbl_w.set(qn("w:w"), str(int(width_pt * 20)))
        tbl_w.set(qn("w:type"), "dxa")

    if indent_pt is not None:
        tbl_ind = tbl_pr.find(qn("w:tblInd"))
        if tbl_ind is None:
            tbl_ind = OxmlElement("w:tblInd")
            tbl_pr.append(tbl_ind)
        tbl_ind.set(qn("w:w"), str(int(max(0.0, indent_pt) * 20)))
        tbl_ind.set(qn("w:type"), "dxa")


def set_cell_margins(cell: Any, margin_pt: float = 2.0) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.find(qn("w:tcMar"))
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for edge in ("top", "start", "bottom", "end"):
        node = tc_mar.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(int(margin_pt * 20)))
        node.set(qn("w:type"), "dxa")


def cluster_numeric(values: list[float], threshold: float) -> list[float]:
    clusters: list[list[float]] = []
    for value in sorted(values):
        if not clusters or abs(median(clusters[-1]) - value) > threshold:
            clusters.append([value])
        else:
            clusters[-1].append(value)
    return [float(median(cluster)) for cluster in clusters]


def choose_consecutive(values: list[float], count: int) -> list[float]:
    if len(values) <= count:
        return values
    best = values[:count]
    best_span = best[-1] - best[0]
    for i in range(1, len(values) - count + 1):
        candidate = values[i : i + count]
        span = candidate[-1] - candidate[0]
        if span > best_span:
            best = candidate
            best_span = span
    return best


def detect_transcript_grid_geometry(
    img: Any,
    page_width_pt: float,
    page_height_pt: float,
) -> dict[str, Any] | None:
    try:
        import cv2
        import numpy as np
    except Exception:
        return None

    arr = np.array(img.convert("RGB"))
    height_px, width_px = arr.shape[:2]
    gray = cv2.cvtColor(arr, cv2.COLOR_RGB2GRAY)
    binary = cv2.adaptiveThreshold(
        gray, 255, cv2.ADAPTIVE_THRESH_MEAN_C, cv2.THRESH_BINARY_INV, 31, 12
    )

    top_offset = int(height_px * 0.28)
    bottom_offset = int(height_px * 0.88)
    left_offset = int(width_px * 0.05)
    right_offset = int(width_px * 0.95)
    roi = binary[top_offset:bottom_offset, left_offset:right_offset]

    h_kernel = cv2.getStructuringElement(cv2.MORPH_RECT, (max(80, width_px // 18), 1))
    v_kernel = cv2.getStructuringElement(cv2.MORPH_RECT, (1, max(50, height_px // 35)))
    horizontal = cv2.morphologyEx(roi, cv2.MORPH_OPEN, h_kernel)
    vertical = cv2.morphologyEx(roi, cv2.MORPH_OPEN, v_kernel)

    xs: list[float] = []
    contours, _ = cv2.findContours(vertical, cv2.RETR_EXTERNAL, cv2.CHAIN_APPROX_SIMPLE)
    for contour in contours:
        x, y, w, h = cv2.boundingRect(contour)
        if h > height_px * 0.10:
            xs.append(left_offset + x + w / 2)
    xs = cluster_numeric(xs, max(10.0, width_px * 0.006))
    xs = choose_consecutive(
        [x for x in xs if width_px * 0.12 <= x <= width_px * 0.94], 6
    )
    if len(xs) < 5:
        return None

    table_left_px = xs[0]
    table_right_px = xs[-1]
    table_width_px = max(1.0, table_right_px - table_left_px)

    ys: list[float] = []
    contours, _ = cv2.findContours(
        horizontal, cv2.RETR_EXTERNAL, cv2.CHAIN_APPROX_SIMPLE
    )
    for contour in contours:
        x, y, w, h = cv2.boundingRect(contour)
        abs_x = left_offset + x
        abs_y = top_offset + y
        center_x = abs_x + w / 2
        if (
            w > table_width_px * 0.42
            and table_left_px - 35 <= center_x <= table_right_px + 35
        ):
            ys.append(abs_y + h / 2)
    ys = cluster_numeric(ys, max(10.0, height_px * 0.004))
    if len(ys) >= 25:
        ys = ys[-25:]
    elif len(ys) < 8:
        return None

    sx = page_width_pt / max(1.0, width_px)
    sy = page_height_pt / max(1.0, height_px)
    x_positions_pt = [x * sx for x in xs]
    y_positions_pt = [y * sy for y in ys]
    col_widths_pt = [
        max(10.0, x_positions_pt[i + 1] - x_positions_pt[i])
        for i in range(len(x_positions_pt) - 1)
    ]
    row_heights_pt = [
        max(9.0, y_positions_pt[i + 1] - y_positions_pt[i])
        for i in range(len(y_positions_pt) - 1)
    ]

    return {
        "xPositionsPx": xs,
        "yPositionsPx": ys,
        "xPositionsPt": x_positions_pt,
        "yPositionsPt": y_positions_pt,
        "colWidthsPt": col_widths_pt,
        "rowHeightsPt": row_heights_pt,
        "tableLeftPt": x_positions_pt[0],
        "tableWidthPt": sum(col_widths_pt),
        "preTableSpacerPt": max(
            0.0, min(120.0, y_positions_pt[0] - page_height_pt * 0.222)
        ),
    }


def words_in_box(
    words: list[WordBox],
    left: float,
    top: float,
    right: float,
    bottom: float,
    min_conf: float = 18,
) -> list[WordBox]:
    out: list[WordBox] = []
    for word in words:
        cx = word.left + word.width / 2
        cy = word.top + word.height / 2
        if left <= cx <= right and top <= cy <= bottom and word.conf >= min_conf:
            out.append(word)
    return out


def cluster_y_centers(values: list[float], threshold: float = 46.0) -> list[float]:
    if not values:
        return []
    clusters: list[list[float]] = []
    for value in sorted(values):
        if not clusters or abs(median(clusters[-1]) - value) > threshold:
            clusters.append([value])
        else:
            clusters[-1].append(value)
    return [median(cluster) for cluster in clusters]


def transcript_row_centers(
    words: list[WordBox], width_px: float, height_px: float
) -> list[float]:
    left = width_px * 0.24
    right = width_px * 0.45
    top = height_px * 0.34
    bottom = height_px * 0.86
    values: list[float] = []
    for word in words:
        token = clean_ocr_token(word.text)
        if word.conf < 40 or not any(ch.isalpha() for ch in token):
            continue
        cx = word.left + word.width / 2
        cy = word.top + word.height / 2
        if left <= cx <= right and top <= cy <= bottom:
            values.append(cy)
    centers = cluster_y_centers(values, 52.0)
    if len(centers) >= 18:
        return centers[:23]

    first = height_px * 0.365
    step = height_px * 0.0214
    return [first + i * step for i in range(23)]
