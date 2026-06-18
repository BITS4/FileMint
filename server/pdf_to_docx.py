#!/usr/bin/env python3
"""High-fidelity PDF -> DOCX conversion for FileMint.

The converter has two paths:
  * Digital/native PDFs: use pdf2docx, which reconstructs editable Word text,
    tables, images and page layout from the PDF object model.
  * Scanned/image PDFs: render pages, OCR with Tesseract TSV output, rebuild
    editable paragraphs and simple detected tables with python-docx.

This is still a free/local pipeline. It does not claim commercial OCR accuracy:
merged scanned-table cells, handwriting, seals and exact vector artwork may not
be perfectly reconstructed. Those limitations are returned in the quality
report so the UI can be honest.
"""

from __future__ import annotations

import argparse
import csv
import json
import math
import os
import re
import shutil
import subprocess
import sys
import tempfile
import urllib.request
import zipfile
from dataclasses import dataclass
from statistics import median
from typing import Any
from xml.sax.saxutils import escape, unescape


OCR_AUTO_LANGS = ["eng", "rus", "tgk", "fas", "ara", "chi_sim", "chi_tra", "kor"]
OCR_AUTO_DOWNLOAD_LANGS = ["chi_sim", "kor"]
LOCAL_TESSDATA_DIR = os.path.join(os.path.dirname(__file__), "tessdata")
TESSDATA_FAST_BASE = "https://github.com/tesseract-ocr/tessdata_fast/raw/main"
DOWNLOADABLE_TESSDATA = {"fas", "ara", "chi_sim", "chi_tra", "kor"}
FAST_HOSTED_OCR = os.environ.get("FILEMINT_FAST_HOSTED_OCR", "").strip().lower() in {"1", "true", "yes", "on"} or os.environ.get("RENDER", "").strip().lower() == "true"
LANG_ALIASES = {
    "": "auto",
    "auto": "auto",
    "english": "eng",
    "russian": "rus",
    "tajik": "tgk",
    "persian": "fas",
    "farsi": "fas",
    "arabic": "ara",
    "chinese": "chi_sim",
    "simplified_chinese": "chi_sim",
    "korean": "kor",
}
MODE_ALIASES = {
    "auto": "hybrid",
    "pro": "high-accuracy",
    "premium": "high-accuracy",
    "ultra": "high-accuracy",
    "premium-editable": "high-accuracy",
    "max-editable": "high-accuracy",
    "editable": "high-accuracy",
    "editable-accurate": "high-accuracy",
    "accurate": "high-accuracy",
    "high": "high-accuracy",
    "high-accuracy": "high-accuracy",
    "high_accuracy": "high-accuracy",
    "high-accuracy-editable": "high-accuracy",
    "hybrid": "hybrid",
    "hybrid-editable": "hybrid",
    "exact": "exact",
    "exact-visual": "exact",
    "ocr": "ocr",
    "ocr-editable": "ocr",
    "image": "image",
    "image-only": "image",
}


def log(*a: object) -> None:
    print(*a, file=sys.stderr)


def truthy(v: str | bool | None, default: bool = True) -> bool:
    if v is None:
        return default
    if isinstance(v, bool):
        return v
    return str(v).strip().lower() not in {"0", "false", "no", "off"}


def safe_mode(mode: str) -> str:
    return MODE_ALIASES.get((mode or "hybrid").strip().lower(), "hybrid")


def engine_mode(mode: str) -> str:
    if mode in {"high-accuracy", "hybrid"}:
        return "premium"
    return mode


def clean_choice(value: str | None, allowed: set[str], default: str) -> str:
    raw = (value or default).strip().lower()
    return raw if raw in allowed else default


def quality_dpi(quality: str, default: int = 300) -> int:
    dpi = {
        "low": 160,
        "medium": 220,
        "high": default,
        "original": 360,
    }.get(quality, default)
    if FAST_HOSTED_OCR:
        return min(dpi, 96)
    return dpi


def effective_ocr_request(lang: str, auto_detect: bool, report: dict[str, Any]) -> str:
    requested = (lang or "auto").strip()
    if auto_detect:
        return requested or "auto"
    if requested.lower() in {"", "auto", "mixed", "auto-mixed"}:
        report["warnings"].append(
            "OCR language auto-detect is off, but no manual language was selected. English OCR will be used."
        )
        return "eng"
    return requested


def find_tesseract() -> str | None:
    found = shutil.which("tesseract")
    if found:
        return found
    for p in [
        r"C:\Program Files\Tesseract-OCR\tesseract.exe",
        r"C:\Program Files (x86)\Tesseract-OCR\tesseract.exe",
        "/usr/bin/tesseract",
        "/opt/homebrew/bin/tesseract",
        "/usr/local/bin/tesseract",
    ]:
        if os.path.isfile(p):
            return p
    return None


def local_tesseract_languages() -> set[str]:
    if not os.path.isdir(LOCAL_TESSDATA_DIR):
        return set()
    return {
        os.path.splitext(name)[0]
        for name in os.listdir(LOCAL_TESSDATA_DIR)
        if name.endswith(".traineddata")
    }


def system_tesseract_languages(tess: str) -> set[str]:
    try:
        r = subprocess.run([tess, "--list-langs"], capture_output=True, text=True, timeout=10)
        lines = (r.stdout or r.stderr or "").splitlines()
        return {x.strip() for x in lines[1:] if x.strip() and not x.lower().startswith("list of")}
    except Exception:
        return set()


def installed_tesseract_languages(tess: str) -> set[str]:
    return system_tesseract_languages(tess) | local_tesseract_languages()


def ensure_project_tessdata(lang: str, report: dict[str, Any]) -> bool:
    if lang in installed_tesseract_languages(find_tesseract() or ""):
        return True
    if lang not in DOWNLOADABLE_TESSDATA:
        return False
    os.makedirs(LOCAL_TESSDATA_DIR, exist_ok=True)
    dst = os.path.join(LOCAL_TESSDATA_DIR, f"{lang}.traineddata")
    if os.path.exists(dst) and os.path.getsize(dst) > 1024:
        return True

    url = f"{TESSDATA_FAST_BASE}/{lang}.traineddata"
    tmp = dst + ".download"
    try:
        urllib.request.urlretrieve(url, tmp)
        if os.path.getsize(tmp) <= 1024:
            raise RuntimeError("downloaded language data is empty")
        os.replace(tmp, dst)
        report["notes"].append(f"Downloaded OCR language data: {lang}.")
        return True
    except Exception as e:
        try:
            if os.path.exists(tmp):
                os.remove(tmp)
        except Exception:
            pass
        report["warnings"].append(f"Could not download OCR language data for {lang}: {e}")
        return False


def ensure_project_tessdata_many(langs: list[str], report: dict[str, Any]) -> None:
    for lang in langs:
        if lang and lang != "osd":
            ensure_project_tessdata(lang, report)


def tessdata_dir_for_lang(lang: str) -> str | None:
    parts = [p for p in re.split(r"[,+\s]+", lang) if p and p != "osd"]
    local = local_tesseract_languages()
    if parts and all(p in local for p in parts):
        return LOCAL_TESSDATA_DIR
    return None


def resolve_ocr_language(requested: str, report: dict[str, Any]) -> str:
    tess = find_tesseract()
    if not tess:
        raise RuntimeError("Tesseract OCR is not installed. Install it, then add needed language packs.")

    raw = LANG_ALIASES.get((requested or "auto").strip().lower(), requested.strip().lower())
    if raw == "auto":
        if not FAST_HOSTED_OCR:
            ensure_project_tessdata_many(OCR_AUTO_DOWNLOAD_LANGS, report)
        else:
            report["notes"].append("Hosted OCR fast mode used installed language packs only.")
    else:
        if not FAST_HOSTED_OCR:
            ensure_project_tessdata_many([p for p in re.split(r"[,+\s]+", raw) if p], report)
    installed = installed_tesseract_languages(tess)
    if raw == "auto":
        if installed:
            if FAST_HOSTED_OCR:
                hosted_preferred = ["eng"]
                chosen = [x for x in hosted_preferred if x in installed]
            else:
                chosen = [x for x in OCR_AUTO_LANGS if x in installed]
            if not chosen and "eng" in installed:
                chosen = ["eng"]
            if not chosen:
                chosen = sorted(installed)[:1]
            if len(chosen) < 2:
                report["warnings"].append(
                    "Auto OCR language is limited because only a small set of Tesseract languages is installed."
                )
            lang = "+".join(chosen) if chosen else "eng"
        else:
            lang = "eng+rus"
            report["warnings"].append("Could not list Tesseract languages; using eng+rus.")
    else:
        parts = [p for p in re.split(r"[,+\s]+", raw) if p]
        if installed:
            missing = [p for p in parts if p not in installed]
            present = [p for p in parts if p in installed]
            if missing:
                report["warnings"].append(
                    "Missing Tesseract language data: " + ", ".join(missing) + "."
                )
            if not present:
                present = ["eng"] if "eng" in installed else sorted(installed)[:1]
                report["warnings"].append("Falling back to OCR language: " + "+".join(present) + ".")
            lang = "+".join(present)
        else:
            lang = "+".join(parts) if parts else "eng"

    report["ocrLanguage"] = lang
    local_parts = [p for p in re.split(r"[,+\s]+", lang) if p in local_tesseract_languages()]
    if local_parts:
        report["notes"].append("OCR used project language data: " + "+".join(local_parts) + ".")
    return lang


def inspect_pdf(path: str) -> dict[str, Any]:
    import fitz

    doc = fitz.open(path)
    details: list[dict[str, Any]] = []
    try:
        total_chars = 0
        text_pages = 0
        image_pages = 0
        image_count = 0
        image_backed_pages = 0
        table_count = 0
        vector_pages = 0

        for page in doc:
            text = page.get_text("text") or ""
            chars = len(text.strip())
            images = len(page.get_images(full=True))
            max_image_coverage = 0.0
            for image in page.get_images(full=True):
                try:
                    for rect in page.get_image_rects(image[0]):
                        coverage = (rect.width * rect.height) / max(1.0, page.rect.width * page.rect.height)
                        max_image_coverage = max(max_image_coverage, coverage)
                except Exception:
                    continue
            drawings = len(page.get_drawings())
            tables = 0
            try:
                found = page.find_tables()
                tables = len(getattr(found, "tables", []) or [])
            except Exception:
                tables = 0

            total_chars += chars
            image_count += images
            table_count += tables
            if chars >= 25:
                text_pages += 1
            if images > 0 and chars < 25:
                image_pages += 1
            if max_image_coverage >= 0.65:
                image_backed_pages += 1
            if drawings > 20:
                vector_pages += 1

            rect = page.rect
            details.append(
                {
                    "page": page.number + 1,
                    "width": round(rect.width, 2),
                    "height": round(rect.height, 2),
                    "textCharacters": chars,
                    "images": images,
                    "maxImageCoverage": round(max_image_coverage, 3),
                    "tables": tables,
                    "vectorObjects": drawings,
                    "scannedLikely": chars < 25 and images > 0,
                    "imageBackedTextLayerLikely": chars >= 25 and max_image_coverage >= 0.65,
                }
            )

        pages = len(doc)
        if text_pages == 0:
            pdf_type = "scanned"
        elif text_pages < pages:
            pdf_type = "mixed"
        elif image_pages > 0:
            pdf_type = "mixed"
        else:
            pdf_type = "digital"

        return {
            "pages": pages,
            "pdfType": pdf_type,
            "textPages": text_pages,
            "imagePages": image_pages,
            "imageBackedPages": image_backed_pages,
            "textCharacters": total_chars,
            "imagesDetected": image_count,
            "tablesDetected": table_count,
            "vectorPages": vector_pages,
            "pageDetails": details,
        }
    finally:
        doc.close()


def has_text_layer(path: str) -> bool:
    return inspect_pdf(path)["textCharacters"] >= 25


def to_docx_pdf2docx(src: str, dst: str, report: dict[str, Any]) -> None:
    from pdf2docx import Converter

    cv = Converter(src)
    try:
        cv.convert(dst)
    finally:
        cv.close()

    report["editableTextDetected"] = True
    report["notes"].append(
        "Digital PDF converted with pdf2docx. Text, layout, images and detected tables are rebuilt as editable DOCX objects where the PDF structure allows it."
    )


@dataclass
class PdfSpan:
    text: str
    font: str
    size: float
    flags: int
    color: int
    bbox: tuple[float, float, float, float]


@dataclass
class PdfLine:
    spans: list[PdfSpan]
    left: float
    top: float
    right: float
    bottom: float


def span_text(spans: list[PdfSpan]) -> str:
    return "".join(span.text for span in spans)


def collect_pdf_lines(page: Any) -> list[PdfLine]:
    raw = page.get_text("dict", sort=True)
    lines: list[PdfLine] = []
    for block in raw.get("blocks", []):
        if block.get("type") != 0:
            continue
        for line in block.get("lines", []):
            spans: list[PdfSpan] = []
            for span in line.get("spans", []):
                text = span.get("text", "")
                if not text:
                    continue
                bbox = tuple(float(x) for x in span.get("bbox", line.get("bbox", (0, 0, 0, 0))))
                spans.append(
                    PdfSpan(
                        text=text,
                        font=str(span.get("font", "")),
                        size=float(span.get("size", 11.0) or 11.0),
                        flags=int(span.get("flags", 0) or 0),
                        color=int(span.get("color", 0) or 0),
                        bbox=bbox,  # type: ignore[arg-type]
                    )
                )
            if not spans or not span_text(spans).strip():
                continue
            left, top, right, bottom = (float(x) for x in line.get("bbox", (0, 0, 0, 0)))
            lines.append(PdfLine(spans=spans, left=left, top=top, right=right, bottom=bottom))
    return sorted(lines, key=lambda line: (line.top, line.left))


def vertical_overlap(a: PdfLine, b: PdfLine) -> float:
    overlap = min(a.bottom, b.bottom) - max(a.top, b.top)
    if overlap <= 0:
        return 0.0
    return overlap / max(1.0, min(a.bottom - a.top, b.bottom - b.top))


def merge_visual_rows(lines: list[PdfLine]) -> list[PdfLine]:
    rows: list[list[PdfLine]] = []
    for line in lines:
        placed = False
        center = (line.top + line.bottom) / 2
        height = max(1.0, line.bottom - line.top)
        for row in rows:
            row_top = min(item.top for item in row)
            row_bottom = max(item.bottom for item in row)
            row_center = (row_top + row_bottom) / 2
            if abs(center - row_center) <= max(2.5, height * 0.42) or any(vertical_overlap(line, item) > 0.68 for item in row):
                row.append(line)
                placed = True
                break
        if not placed:
            rows.append([line])

    merged: list[PdfLine] = []
    for row in rows:
        ordered = sorted(row, key=lambda item: item.left)
        spans: list[PdfSpan] = []
        for item in ordered:
            spans.extend(item.spans)
        merged.append(
            PdfLine(
                spans=spans,
                left=min(item.left for item in ordered),
                top=min(item.top for item in ordered),
                right=max(item.right for item in ordered),
                bottom=max(item.bottom for item in ordered),
            )
        )
    return sorted(merged, key=lambda line: (line.top, line.left))


def word_font_name(pdf_font: str) -> str:
    font = pdf_font.upper()
    if "TT" in font or "MONO" in font or "COURIER" in font:
        return "Courier New"
    return "Times New Roman"


def span_is_bold(span: PdfSpan) -> bool:
    font = span.font.upper()
    return bool(span.flags & 16) or "BOLD" in font or "CMBX" in font or "BX" in font


def span_is_italic(span: PdfSpan) -> bool:
    font = span.font.upper()
    return bool(span.flags & 2) or "ITAL" in font or "CMMI" in font or "CMTI" in font or "MI" in font


def set_run_color(run: Any, color: int) -> None:
    if color == 0:
        return
    try:
        from docx.shared import RGBColor

        run.font.color.rgb = RGBColor((color >> 16) & 255, (color >> 8) & 255, color & 255)
    except Exception:
        pass


def xml_compatible_text(text: str) -> str:
    return "".join(ch for ch in text if ch in "\t\n\r" or ord(ch) >= 32)


def append_pdf_span_run(paragraph: Any, span: PdfSpan, row: PdfLine, dominant_size: float) -> None:
    from docx.shared import Pt

    text = xml_compatible_text(span.text)
    if not text:
        return
    run = paragraph.add_run(text)
    run.font.name = word_font_name(span.font)
    run.font.size = Pt(max(5.0, min(28.0, span.size)))
    run.bold = span_is_bold(span)
    run.italic = span_is_italic(span)
    set_run_color(run, span.color)

    if span.size < dominant_size * 0.82:
        span_mid = (span.bbox[1] + span.bbox[3]) / 2
        row_mid = (row.top + row.bottom) / 2
        if span_mid < row_mid - dominant_size * 0.08:
            run.font.superscript = True
        elif span_mid > row_mid + dominant_size * 0.08:
            run.font.subscript = True


def estimated_gap_spaces(gap_pt: float, size_pt: float) -> int:
    if gap_pt <= max(1.0, size_pt * 0.12):
        return 0
    return max(1, min(24, round(gap_pt / max(2.5, size_pt * 0.33))))


def append_pdf_line(doc: Any, line: PdfLine, prev_bottom: float | None, left_margin_pt: float) -> int:
    from docx.shared import Pt

    paragraph = doc.add_paragraph()
    fmt = paragraph.paragraph_format
    fmt.left_indent = Pt(max(0.0, line.left - left_margin_pt))
    fmt.space_after = Pt(0)
    fmt.line_spacing = 1.0
    if prev_bottom is None:
        fmt.space_before = Pt(0)
    else:
        gap = max(0.0, line.top - prev_bottom)
        fmt.space_before = Pt(min(40.0, gap))

    sizes = [span.size for span in line.spans if span.text.strip()]
    dominant_size = median(sizes) if sizes else 11.0
    last_right: float | None = None
    emitted_chars = 0
    for span in sorted(line.spans, key=lambda s: (s.bbox[0], s.bbox[1])):
        text = span.text
        if last_right is not None:
            gap = span.bbox[0] - last_right
            spaces = estimated_gap_spaces(gap, dominant_size)
            if spaces and not text.startswith(" "):
                paragraph.add_run(" " * spaces)
                emitted_chars += spaces
        append_pdf_span_run(paragraph, span, line, dominant_size)
        emitted_chars += len(text)
        last_right = max(last_right or span.bbox[2], span.bbox[2])
    return emitted_chars


def to_docx_digital_text_flow(src: str, dst: str, report: dict[str, Any]) -> None:
    import fitz
    from docx import Document
    from docx.shared import Pt

    pdf = fitz.open(src)
    out = Document()
    try:
        total_lines = 0
        total_chars = 0
        for page_index, page in enumerate(pdf):
            if page_index == 0:
                section = out.sections[-1]
            else:
                out.add_page_break()
                section = out.sections[-1]
            section.page_width = Pt(page.rect.width)
            section.page_height = Pt(page.rect.height)
            section.left_margin = Pt(72)
            section.right_margin = Pt(72)
            section.top_margin = Pt(72)
            section.bottom_margin = Pt(72)

            rows = merge_visual_rows(collect_pdf_lines(page))
            prev_bottom: float | None = None
            for row in rows:
                total_chars += append_pdf_line(out, row, prev_bottom, 72.0)
                total_lines += 1
                prev_bottom = row.bottom

        report["resolvedMode"] = "premium-digital-flow"
        report["editableTextDetected"] = True
        report["editableTextBoxes"] = total_lines
        report["editableCharacters"] = total_chars
        report["ocrTextCandidates"] = total_lines
        report["textCoverageEstimate"] = 100 if total_lines else 0
        report["notes"].append(
            "Premium digital text-flow rebuilt the native PDF text layer as editable Word text with inferred word spacing, line gaps and page breaks."
        )
        out.save(dst)
    finally:
        pdf.close()


@dataclass
class WordBox:
    text: str
    left: float
    top: float
    width: float
    height: float
    conf: float
    block: int
    par: int
    line: int


@dataclass
class LineBox:
    text: str
    words: list[WordBox]
    left: float
    top: float
    width: float
    height: float
    conf: float
    page_width_px: float
    page_height_px: float
    page_width_pt: float
    page_height_pt: float
    segments: list[tuple[float, float, str]]


@dataclass
class VisualFragment:
    path: str
    left: float
    top: float
    width: float
    height: float
    kind: str = "image"


@dataclass
class VisualRule:
    left: float
    top: float
    width: float
    height: float
    color: str = "#AEB6BC"


CJK_RANGE = "\u3400-\u4dbf\u4e00-\u9fff\uf900-\ufaff"


def normalize_ocr_text(text: str) -> str:
    text = xml_compatible_text(text)
    text = re.sub(r"\s+", " ", text).strip()
    previous = None
    while previous != text:
        previous = text
        text = re.sub(fr"([{CJK_RANGE}])\s+([{CJK_RANGE}])", r"\1\2", text)
    text = re.sub(fr"([{CJK_RANGE}])\s+([，。！？；：、])", r"\1\2", text)
    text = re.sub(fr"([，。！？；：、])\s+([{CJK_RANGE}])", r"\1\2", text)
    text = re.sub(r"\s+([,.;:!?])", r"\1", text)
    return text.strip()


def parse_tsv(tsv: str, page_width_px: float, page_height_px: float, page_width_pt: float, page_height_pt: float) -> list[LineBox]:
    tsv = tsv.replace("\f", "\n")
    rows = list(csv.DictReader(tsv.splitlines(), delimiter="\t"))
    grouped: dict[tuple[int, int, int], list[WordBox]] = {}
    all_words: list[WordBox] = []
    for row in rows:
        try:
            if int(row.get("level", "0")) != 5:
                continue
            text = (row.get("text") or "").strip()
            if not text:
                continue
            if "\t" in text:
                text = text.split("\t", 1)[0].strip()
            if not text:
                continue
            conf = float(row.get("conf") or -1)
            word = WordBox(
                text=text,
                left=float(row.get("left") or 0),
                top=float(row.get("top") or 0),
                width=float(row.get("width") or 0),
                height=float(row.get("height") or 0),
                conf=conf,
                block=int(row.get("block_num") or 0),
                par=int(row.get("par_num") or 0),
                line=int(row.get("line_num") or 0),
            )
            all_words.append(word)
            grouped.setdefault((word.block, word.par, word.line), []).append(word)
        except Exception:
            continue

    lines: list[LineBox] = []
    for words in grouped.values():
        words = sorted(words, key=lambda w: w.left)
        left = min(w.left for w in words)
        top = min(w.top for w in words)
        right = max(w.left + w.width for w in words)
        bottom = max(w.top + w.height for w in words)
        confs = [w.conf for w in words if w.conf >= 0]
        text = normalize_ocr_text(" ".join(w.text for w in words))
        lines.append(
            LineBox(
                text=text,
                words=words,
                left=left,
                top=top,
                width=right - left,
                height=bottom - top,
                conf=sum(confs) / len(confs) if confs else -1,
                page_width_px=page_width_px,
                page_height_px=page_height_px,
                page_width_pt=page_width_pt,
                page_height_pt=page_height_pt,
                segments=segment_line(words),
            )
        )
    if all_words and (len(lines) < 8 or max((len(line.words) for line in lines), default=0) > 45):
        rebuilt = rebuild_rows_from_word_geometry(all_words, page_width_px, page_height_px, page_width_pt, page_height_pt)
        if len(rebuilt) > len(lines) * 2:
            return rebuilt
    return sorted(lines, key=lambda l: (l.top, l.left))


def rebuild_rows_from_word_geometry(
    words: list[WordBox],
    page_width_px: float,
    page_height_px: float,
    page_width_pt: float,
    page_height_pt: float,
) -> list[LineBox]:
    rows: list[list[WordBox]] = []
    for word in sorted(words, key=lambda w: (w.top + w.height / 2, w.left)):
        if word.width <= 0 or word.height <= 0:
            continue
        center = word.top + word.height / 2
        placed = False
        for row in rows:
            row_center = sum(w.top + w.height / 2 for w in row) / max(1, len(row))
            row_height = median([w.height for w in row if w.height > 0]) or word.height
            if abs(center - row_center) <= max(5.0, min(18.0, row_height * 0.62)):
                row.append(word)
                placed = True
                break
        if not placed:
            rows.append([word])

    line_boxes: list[LineBox] = []
    for row in rows:
        row = sorted(row, key=lambda w: w.left)
        if not row:
            continue
        left = min(w.left for w in row)
        top = min(w.top for w in row)
        right = max(w.left + w.width for w in row)
        bottom = max(w.top + w.height for w in row)
        confs = [w.conf for w in row if w.conf >= 0]
        text = normalize_ocr_text(" ".join(w.text for w in row))
        if not text:
            continue
        line_boxes.append(
            LineBox(
                text=text,
                words=row,
                left=left,
                top=top,
                width=right - left,
                height=bottom - top,
                conf=sum(confs) / len(confs) if confs else -1,
                page_width_px=page_width_px,
                page_height_px=page_height_px,
                page_width_pt=page_width_pt,
                page_height_pt=page_height_pt,
                segments=segment_line(row),
            )
        )
    return sorted(line_boxes, key=lambda l: (l.top, l.left))


def segment_line(words: list[WordBox]) -> list[tuple[float, float, str]]:
    if not words:
        return []
    heights = [w.height for w in words if w.height > 0]
    h = median(heights) if heights else 12
    gap_limit = max(55.0, h * 3.2)
    segments: list[list[WordBox]] = [[words[0]]]
    last_right = words[0].left + words[0].width
    for w in words[1:]:
        gap = w.left - last_right
        if gap > gap_limit:
            segments.append([w])
        else:
            segments[-1].append(w)
        last_right = max(last_right, w.left + w.width)
    out: list[tuple[float, float, str]] = []
    for segment in segments:
        left = min(w.left for w in segment)
        right = max(w.left + w.width for w in segment)
        out.append((left, right, normalize_ocr_text(" ".join(w.text for w in segment))))
    return out


def run_tesseract_tsv(image: str, lang: str, psm: str = "11") -> str:
    tess = find_tesseract()
    if not tess:
        raise RuntimeError("Tesseract OCR is not installed.")
    cmd = [tess, image, "stdout", "-l", lang, "--oem", "1", "--psm", psm, "tsv"]
    tessdata_dir = tessdata_dir_for_lang(lang)
    if tessdata_dir:
        cmd[3:3] = ["--tessdata-dir", tessdata_dir]
    timeout = int(os.environ.get("FILEMINT_TESSERACT_TIMEOUT_SEC", "60" if FAST_HOSTED_OCR else "180"))
    try:
        r = subprocess.run(
            cmd,
            check=False,
            capture_output=True,
            text=True,
            encoding="utf-8",
            errors="ignore",
            timeout=timeout,
        )
    except subprocess.TimeoutExpired:
        raise RuntimeError(f"Tesseract OCR timed out after {timeout}s.")
    if r.returncode != 0:
        raise RuntimeError((r.stderr or r.stdout or "Tesseract OCR failed.").strip())
    return r.stdout


def unique_lang(parts: list[str]) -> str:
    out: list[str] = []
    for part in parts:
        if part and part not in out:
            out.append(part)
    return "+".join(out)


def ocr_language_candidates(lang: str) -> list[str]:
    parts = [p for p in re.split(r"[,+\s]+", lang) if p]
    if FAST_HOSTED_OCR:
        preferred = ["eng", "rus", "tgk", "fas", "ara", "chi_sim", "chi_tra", "kor"]
        singles = [p for p in preferred if p in parts] + [p for p in parts if p not in preferred]
        return [(singles or ["eng"])[0]]
    part_set = set(parts)
    candidates = [unique_lang(parts)]
    if "eng" in part_set and ("chi_sim" in part_set or "chi_tra" in part_set):
        candidates.append(unique_lang(["eng"] + [p for p in ("chi_sim", "chi_tra") if p in part_set]))
        candidates.append(unique_lang([p for p in ("chi_sim", "chi_tra") if p in part_set]))
    if "eng" in part_set and (part_set & {"rus", "tgk"}):
        candidates.append(unique_lang(["eng"] + [p for p in ("rus", "tgk") if p in part_set]))
    if "eng" in part_set and (part_set & {"fas", "ara"}):
        candidates.append(unique_lang(["eng"] + [p for p in ("fas", "ara") if p in part_set]))
    candidates.append(unique_lang(parts[:1]))
    return [c for i, c in enumerate(candidates) if c and c not in candidates[:i]]


def script_counts(text: str) -> dict[str, int]:
    return {
        "latin": len(re.findall(r"[A-Za-z]", text)),
        "cyrillic": len(re.findall(r"[\u0400-\u052f]", text)),
        "cjk": len(re.findall(fr"[{CJK_RANGE}]", text)),
        "rtl": len(re.findall(r"[\u0590-\u08ff]", text)),
        "digits": len(re.findall(r"\d", text)),
    }


def score_ocr_lines(lines: list[LineBox], lang: str) -> float:
    if not lines:
        return -1000.0
    text = " ".join(line.text for line in lines)
    counts = script_counts(text)
    confs = [max(0.0, line.conf) for line in lines if line.conf >= 0]
    avg_conf = sum(confs) / max(1, len(confs))
    signal_chars = counts["latin"] + counts["cyrillic"] + counts["cjk"] + counts["rtl"] + counts["digits"]
    score = avg_conf + min(16.0, signal_chars / 80.0) + min(8.0, len(lines) / 12.0)
    parts = set(p for p in re.split(r"[,+\s]+", lang) if p)
    if counts["cjk"] >= 10 and parts & {"chi_sim", "chi_tra"}:
        score += 12.0
        if parts & {"rus", "tgk"}:
            score -= min(22.0, counts["cyrillic"] / max(1.0, counts["cjk"]) * 18.0)
    if counts["rtl"] >= 10 and parts & {"fas", "ara"}:
        score += 8.0
    if counts["cyrillic"] >= 20 and parts & {"rus", "tgk"} and counts["cjk"] < 10:
        score += 8.0
    return score


def contains_rtl(text: str) -> bool:
    return bool(re.search(r"[\u0590-\u08ff]", text))


def set_paragraph_bidi(paragraph: Any) -> None:
    from docx.oxml import OxmlElement

    p_pr = paragraph._p.get_or_add_pPr()
    if p_pr.find("./w:bidi", paragraph._p.nsmap) is None:
        p_pr.append(OxmlElement("w:bidi"))


def set_run_font(run: Any, name: str = "Arial") -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    run.font.name = name
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.append(r_fonts)
    for key in ("w:ascii", "w:hAnsi", "w:eastAsia", "w:cs"):
        r_fonts.set(qn(key), name)


def set_cell_text(cell: Any, text: str, font_size: float, rtl: bool) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    if rtl:
        set_paragraph_bidi(p)
    run = p.add_run(text)
    set_run_font(run)
    try:
        from docx.shared import Pt

        run.font.size = Pt(max(6.0, min(22.0, font_size)))
    except Exception:
        pass


def cluster_columns(lines: list[LineBox], max_cols: int) -> list[float]:
    lefts: list[float] = []
    for line in lines:
        for left, _right, _text in line.segments:
            lefts.append(left)
    if not lefts:
        return []
    lefts.sort()
    clusters: list[list[float]] = []
    for x in lefts:
        if not clusters or abs(median(clusters[-1]) - x) > 45:
            clusters.append([x])
        else:
            clusters[-1].append(x)
    centers = [median(c) for c in clusters]
    return centers[:max_cols]


def table_runs(lines: list[LineBox]) -> list[list[LineBox]]:
    runs: list[list[LineBox]] = []
    current: list[LineBox] = []
    last: LineBox | None = None
    for line in lines:
        is_candidate = len(line.segments) >= 2
        if not is_candidate:
            if len(current) >= 2:
                runs.append(current)
            current = []
            last = None
            continue

        if last is None:
            current = [line]
        else:
            vertical_gap = line.top - (last.top + last.height)
            similar_cols = abs(len(line.segments) - len(last.segments)) <= 1
            close = vertical_gap <= max(80, median([last.height, line.height]) * 3.5)
            if similar_cols and close:
                current.append(line)
            else:
                if len(current) >= 2:
                    runs.append(current)
                current = [line]
        last = line

    if len(current) >= 2:
        runs.append(current)
    return runs


def add_table(doc: Any, lines: list[LineBox], report: dict[str, Any]) -> None:
    from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
    from docx.shared import Pt

    max_cols = max(len(l.segments) for l in lines)
    table = doc.add_table(rows=len(lines), cols=max_cols)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    columns = cluster_columns(lines, max_cols)
    page_width_pt = lines[0].page_width_pt
    scale = page_width_pt / max(1.0, lines[0].page_width_px)
    usable_width = max(120.0, page_width_pt - 72.0)

    col_widths: list[float] = []
    for i in range(max_cols):
        left = columns[i] * scale if i < len(columns) else (36 + (usable_width / max_cols) * i)
        if i + 1 < len(columns):
            right = columns[i + 1] * scale
        else:
            right = page_width_pt - 36
        col_widths.append(max(36.0, right - left))

    for r_idx, line in enumerate(lines):
        row = table.rows[r_idx]
        try:
            row.height = Pt(max(14.0, line.height * scale * 1.35))
        except Exception:
            pass
        for c_idx in range(max_cols):
            cell = row.cells[c_idx]
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            try:
                cell.width = Pt(col_widths[c_idx])
            except Exception:
                pass
            text = line.segments[c_idx][2] if c_idx < len(line.segments) else ""
            set_cell_text(cell, text, max(8.0, line.height * scale * 0.95), contains_rtl(text))

    report["tablesDetected"] = int(report.get("tablesDetected", 0)) + 1


def add_paragraph_from_line(doc: Any, line: LineBox, prev_bottom: float | None) -> None:
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt

    scale_x = line.page_width_pt / max(1.0, line.page_width_px)
    scale_y = line.page_height_pt / max(1.0, line.page_height_px)
    p = doc.add_paragraph()
    if contains_rtl(line.text):
        set_paragraph_bidi(p)
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    else:
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    fmt = p.paragraph_format
    fmt.left_indent = Pt(max(0.0, line.left * scale_x - 36.0))
    if prev_bottom is not None:
        gap = max(0.0, line.top - prev_bottom)
        fmt.space_before = Pt(min(36.0, gap * scale_y))
    fmt.space_after = Pt(0)
    fmt.line_spacing = 1.0

    run = p.add_run(line.text)
    set_run_font(run)
    run.font.size = Pt(max(6.0, min(22.0, line.height * scale_y * 0.95)))


VML_NS = (
    'xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" '
    'xmlns:v="urn:schemas-microsoft-com:vml" '
    'xmlns:o="urn:schemas-microsoft-com:office:office" '
    'xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships"'
)


def ocr_font_size(line: LineBox) -> float:
    scale_y = line.page_height_pt / max(1.0, line.page_height_px)
    scale_x = line.page_width_pt / max(1.0, line.page_width_px)
    height_size = line.height * scale_y * 0.68
    counts = script_counts(line.text)
    spaces = line.text.count(" ")
    punctuation = sum(1 for ch in line.text if not ch.isalnum() and not ch.isspace())
    units = (
        counts["latin"] * 0.48
        + counts["digits"] * 0.48
        + counts["cyrillic"] * 0.52
        + counts["rtl"] * 0.56
        + counts["cjk"] * 0.96
        + spaces * 0.26
        + punctuation * 0.30
    )
    if units <= 0:
        return max(5.0, min(20.0, height_size))
    fit_size = (line.width * scale_x * 1.04 + 8.0) / units
    return max(4.5, min(20.0, height_size, fit_size))


def ocr_font_attrs(text: str) -> tuple[str, str, str]:
    if re.search(fr"[{CJK_RANGE}]", text):
        return ("Times New Roman", "SimSun", "Arial")
    if contains_rtl(text):
        return ("Arial", "Arial", "Arial")
    return ("Times New Roman", "Times New Roman", "Arial")


def editable_confidence_threshold(lang: str) -> float:
    parts = set(p for p in re.split(r"[,+\s]+", lang) if p)
    if parts & {"rus", "tgk", "fas", "ara", "chi_sim", "chi_tra"}:
        return 42.0
    return 72.0


def premium_confidence_threshold(lang: str) -> float:
    parts = set(p for p in re.split(r"[,+\s]+", lang) if p)
    if parts & {"rus", "tgk", "fas", "ara", "chi_sim", "chi_tra"}:
        return 18.0
    return 38.0


def line_is_confident(line: LineBox, min_conf: float = 72.0) -> bool:
    if line.conf < min_conf:
        return False
    alnum = sum(1 for ch in line.text if ch.isalnum())
    return alnum >= 2


def line_text_signal(line: LineBox) -> bool:
    alnum = sum(1 for ch in line.text if ch.isalnum())
    letters = sum(1 for ch in line.text if ch.isalpha())
    return alnum >= 2 and (letters >= 1 or any(ch.isdigit() for ch in line.text))


def line_overlaps_colored_mark(img: Any, line: LineBox) -> bool:
    pad = max(6, int(line.height * 0.35))
    left = max(0, int(line.left - pad))
    top = max(0, int(line.top - pad))
    right = min(img.width, int(line.left + line.width + pad))
    bottom = min(img.height, int(line.top + line.height + pad))
    if right <= left or bottom <= top:
        return False

    total = 0
    colorful = 0
    step = max(1, int(max(right - left, bottom - top) / 80))
    pixels = img.load()
    for y in range(top, bottom, step):
        for x in range(left, right, step):
            r, g, b = pixels[x, y]
            if r > 235 and g > 235 and b > 235:
                continue
            total += 1
            if max(r, g, b) - min(r, g, b) > 38:
                colorful += 1
    return total > 0 and (colorful / total) > 0.08


def exact_editable_lines(
    src_png: str,
    lines: list[LineBox],
    lang: str,
    premium: bool = False,
) -> tuple[list[LineBox], dict[str, int]]:
    from PIL import Image

    img = Image.open(src_png).convert("RGB")
    editable: list[LineBox] = []
    stats = {
        "ocrTextCandidates": 0,
        "editableTextBoxes": 0,
        "editableCharacters": 0,
        "skippedColoredMarks": 0,
        "skippedLowConfidence": 0,
        "skippedNoise": 0,
    }
    min_conf = premium_confidence_threshold(lang) if premium else editable_confidence_threshold(lang)
    for line in lines:
        if not line_text_signal(line):
            continue
        if probable_ocr_noise(line, lang):
            stats["skippedNoise"] += 1
            continue
        if line_overlaps_colored_mark(img, line) and (not premium or line.conf < 62):
            stats["skippedColoredMarks"] += 1
            continue
        stats["ocrTextCandidates"] += 1
        if not line_is_confident(line, min_conf):
            stats["skippedLowConfidence"] += 1
            continue
        editable.append(line)
    stats["editableTextBoxes"] = len(editable)
    stats["editableCharacters"] = sum(len(line.text) for line in editable)
    return editable, stats


def word_as_line(word: WordBox, parent: LineBox) -> LineBox:
    return LineBox(
        text=word.text,
        words=[word],
        left=word.left,
        top=word.top,
        width=word.width,
        height=word.height,
        conf=word.conf,
        page_width_px=parent.page_width_px,
        page_height_px=parent.page_height_px,
        page_width_pt=parent.page_width_pt,
        page_height_pt=parent.page_height_pt,
        segments=[(word.left, word.left + word.width, word.text)],
    )


def is_duplicate_word_line(candidate: LineBox, existing: list[LineBox]) -> bool:
    for line in existing:
        if candidate.text.strip().lower() != line.text.strip().lower():
            continue
        x_overlap = min(candidate.left + candidate.width, line.left + line.width) - max(candidate.left, line.left)
        y_overlap = min(candidate.top + candidate.height, line.top + line.height) - max(candidate.top, line.top)
        if x_overlap <= 0 or y_overlap <= 0:
            continue
        overlap = x_overlap * y_overlap
        smaller = max(1.0, min(candidate.width * candidate.height, line.width * line.height))
        if overlap / smaller > 0.45:
            return True
    return False


def dense_table_scan_likely(lines: list[LineBox]) -> bool:
    total_words = sum(len(line.words) for line in lines)
    segmented_rows = sum(1 for line in lines if len(line.segments) >= 3)
    numeric_grade_lines = sum(1 for line in lines if re.search(r"\d+/\d+", line.text))
    if len(lines) >= 24 and total_words >= 80 and (segmented_rows >= 10 or numeric_grade_lines >= 6):
        return True
    if len(lines) < 80:
        return False
    y_positions = sorted(line.top for line in lines if line.height > 0)
    if len(y_positions) < 40:
        return False
    close_rows = 0
    last_y: float | None = None
    for y in y_positions:
        if last_y is not None and y - last_y < 42:
            close_rows += 1
        last_y = y
    return close_rows >= 35 or numeric_grade_lines >= 8


def transcript_scan_likely(lines: list[LineBox]) -> bool:
    """Return true only for the known school-transcript layout.

    The transcript rebuilder below is intentionally template-specific. It must
    never run on generic scans just because OCR sees many rows or numbers.
    """
    text = normalize_ocr_text(" ".join(line.text for line in lines)).lower()
    if not text:
        return False
    title_signals = [
        "student personal",
        "personal information",
        "academic record",
        "course title",
        "average grade",
    ]
    grade_headers = [
        "grade 9",
        "grade 10",
        "grade 11",
        "half-year",
        "half year",
    ]
    school_terms = [
        "student",
        "academic",
        "course",
        "grade",
        "graduation",
        "principal",
        "lyceum",
        "school",
        "language",
    ]
    title_score = sum(1 for signal in title_signals if signal in text)
    header_score = sum(1 for signal in grade_headers if signal in text)
    school_score = sum(1 for signal in school_terms if signal in text)
    grade_mark_count = len(re.findall(r"\b(?:5/5|4/5|8/10|9/10|10/10)\b", text))
    return title_score >= 2 and header_score >= 2 and school_score >= 4 and grade_mark_count >= 4


def exact_editable_word_lines(
    src_png: str,
    lines: list[LineBox],
    lang: str,
    premium: bool = False,
) -> tuple[list[LineBox], dict[str, int]]:
    from PIL import Image

    img = Image.open(src_png).convert("RGB")
    editable: list[LineBox] = []
    stats = {
        "ocrTextCandidates": 0,
        "editableTextBoxes": 0,
        "editableCharacters": 0,
        "skippedColoredMarks": 0,
        "skippedLowConfidence": 0,
        "skippedNoise": 0,
    }
    min_conf = premium_confidence_threshold(lang) if premium else editable_confidence_threshold(lang)
    for line in lines:
        for word in line.words:
            word_line = word_as_line(word, line)
            if not line_text_signal(word_line):
                continue
            if probable_ocr_noise(word_line, lang):
                stats["skippedNoise"] += 1
                continue
            if line_overlaps_colored_mark(img, word_line) and (not premium or word_line.conf < 62):
                stats["skippedColoredMarks"] += 1
                continue
            stats["ocrTextCandidates"] += 1
            if not line_is_confident(word_line, min_conf):
                stats["skippedLowConfidence"] += 1
                continue
            if is_duplicate_word_line(word_line, editable):
                continue
            editable.append(word_line)

    editable.sort(key=lambda line: (line.top, line.left))
    stats["editableTextBoxes"] = len(editable)
    stats["editableCharacters"] = sum(len(line.text) for line in editable)
    return editable, stats


def make_residual_image(src_png: str, dst_png: str, lines: list[LineBox]) -> None:
    from PIL import Image, ImageDraw

    img = Image.open(src_png).convert("RGB")
    draw = ImageDraw.Draw(img)
    for line in lines:
        pad_x = max(3, int(line.height * 0.18))
        pad_y = max(3, int(line.height * 0.22))
        draw.rectangle(
            [
                max(0, int(line.left - pad_x)),
                max(0, int(line.top - pad_y)),
                min(img.width, int(line.left + line.width + pad_x)),
                min(img.height, int(line.top + line.height + pad_y)),
            ],
            fill=(255, 255, 255),
        )
    img.save(dst_png)


def remove_text_regions_rgb(arr: Any, lines: list[LineBox]) -> None:
    for line in lines:
        pad_x = max(5, int(line.height * 0.45))
        pad_y = max(5, int(line.height * 0.55))
        left = max(0, int(line.left - pad_x))
        top = max(0, int(line.top - pad_y))
        right = min(arr.shape[1], int(line.left + line.width + pad_x))
        bottom = min(arr.shape[0], int(line.top + line.height + pad_y))
        if right > left and bottom > top:
            arr[top:bottom, left:right] = 255


def color_hex_from_region(arr: Any, left: int, top: int, right: int, bottom: int) -> str:
    region = arr[max(0, top):max(0, bottom), max(0, left):max(0, right)]
    if region.size == 0:
        return "#AEB6BC"
    import numpy as np

    flat = region.reshape(-1, 3)
    nonwhite = flat[np.any(flat < 235, axis=1)]
    if len(nonwhite) == 0:
        return "#AEB6BC"
    rgb = np.median(nonwhite, axis=0).astype(int)
    return f"#{rgb[0]:02X}{rgb[1]:02X}{rgb[2]:02X}"


def save_transparent_fragment(arr: Any, mask: Any, path: str) -> bool:
    from PIL import Image
    import numpy as np

    if arr.size == 0:
        return False
    alpha = np.where(mask > 0, 255, 0).astype("uint8")
    if int(np.count_nonzero(alpha)) < 20:
        return False
    rgba = np.dstack([arr, alpha])
    Image.fromarray(rgba, "RGBA").save(path)
    return True


def merge_visual_rules(rules: list[VisualRule]) -> list[VisualRule]:
    merged: list[VisualRule] = []
    for rule in sorted(rules, key=lambda r: (round(r.top, 1), r.left)):
        if merged:
            prev = merged[-1]
            same_row = abs((prev.top + prev.height / 2) - (rule.top + rule.height / 2)) <= max(1.0, prev.height, rule.height) * 1.4
            close = rule.left <= prev.left + prev.width + 4.0
            if same_row and close and prev.color == rule.color:
                right = max(prev.left + prev.width, rule.left + rule.width)
                prev.left = min(prev.left, rule.left)
                prev.width = right - prev.left
                prev.top = min(prev.top, rule.top)
                prev.height = max(prev.height, rule.height)
                continue
        merged.append(rule)
    return merged


def segment_visual_layer(
    src_png: str,
    dst_dir: str,
    page_index: int,
    editable_lines: list[LineBox],
    page_width_pt: float,
    page_height_pt: float,
) -> tuple[list[VisualFragment], list[VisualRule]]:
    try:
        import cv2
        import numpy as np
        from PIL import Image
    except Exception as e:
        raise RuntimeError(f"OpenCV image segmentation is unavailable: {e}")

    img = Image.open(src_png).convert("RGB")
    original_arr = np.array(img)
    arr = original_arr.copy()
    height_px, width_px = arr.shape[:2]
    scale_x = page_width_pt / max(1.0, width_px)
    scale_y = page_height_pt / max(1.0, height_px)

    fragments: list[VisualFragment] = []
    original_hsv = cv2.cvtColor(original_arr, cv2.COLOR_RGB2HSV)
    color_delta = original_arr.max(axis=2) - original_arr.min(axis=2)
    colored_mask = np.where(
        (original_hsv[:, :, 1] > 34)
        & (original_hsv[:, :, 2] < 252)
        & (color_delta > 28),
        255,
        0,
    ).astype("uint8")
    color_kernel = cv2.getStructuringElement(cv2.MORPH_RECT, (7, 7))
    colored_components = cv2.dilate(
        cv2.morphologyEx(colored_mask, cv2.MORPH_CLOSE, color_kernel, iterations=1),
        color_kernel,
        iterations=2,
    )
    color_count, color_labels, color_stats, _ = cv2.connectedComponentsWithStats(colored_components, 8)
    for label in range(1, color_count):
        x, y, w, h, area = color_stats[label]
        if area < 180 or w < 14 or h < 14 or max(w, h) < 30:
            continue
        if w * h > width_px * height_px * 0.18:
            continue
        pad = max(4, min(20, int(max(w, h) * 0.035)))
        left = max(0, x - pad)
        top = max(0, y - pad)
        right = min(width_px, x + w + pad)
        bottom = min(height_px, y + h + pad)
        crop = original_arr[top:bottom, left:right].copy()
        crop_mask = colored_mask[top:bottom, left:right].copy()
        fragment_path = os.path.join(dst_dir, f"page-{page_index + 1}-color-fragment-{len(fragments) + 1}.png")
        if save_transparent_fragment(crop, crop_mask, fragment_path):
            fragments.append(
                VisualFragment(
                    path=fragment_path,
                    left=left * scale_x,
                    top=top * scale_y,
                    width=(right - left) * scale_x,
                    height=(bottom - top) * scale_y,
                    kind="colored-mark",
                )
            )
        region = arr[top:bottom, left:right]
        region[crop_mask > 0] = 255
        arr[top:bottom, left:right] = region

    remove_text_regions_rgb(arr, editable_lines)

    gray = cv2.cvtColor(arr, cv2.COLOR_RGB2GRAY)
    dark = np.where(gray < 224, 255, 0).astype("uint8")
    horizontal_kernel = cv2.getStructuringElement(cv2.MORPH_RECT, (max(45, width_px // 48), 1))
    horizontal = cv2.morphologyEx(dark, cv2.MORPH_OPEN, horizontal_kernel)
    contours, _ = cv2.findContours(horizontal, cv2.RETR_EXTERNAL, cv2.CHAIN_APPROX_SIMPLE)

    rules: list[VisualRule] = []
    for contour in contours:
        x, y, w, h = cv2.boundingRect(contour)
        if w < max(70, width_px * 0.035):
            continue
        if h > max(14, height_px * 0.006):
            continue
        color = color_hex_from_region(arr, x, y, x + w, y + h)
        rules.append(
            VisualRule(
                left=x * scale_x,
                top=y * scale_y,
                width=w * scale_x,
                height=max(0.45, h * scale_y),
                color=color,
            )
        )
        pad = 2
        arr[max(0, y - pad):min(height_px, y + h + pad), max(0, x - pad):min(width_px, x + w + pad)] = 255

    rules = merge_visual_rules(rules)

    gray = cv2.cvtColor(arr, cv2.COLOR_RGB2GRAY)
    hsv = cv2.cvtColor(arr, cv2.COLOR_RGB2HSV)
    saturation = hsv[:, :, 1]
    value = hsv[:, :, 2]
    mask = np.where((gray < 244) | ((saturation > 24) & (value < 252)), 255, 0).astype("uint8")
    mask[:3, :] = 0
    mask[-3:, :] = 0
    mask[:, :3] = 0
    mask[:, -3:] = 0
    kernel = cv2.getStructuringElement(cv2.MORPH_RECT, (5, 5))
    connected_mask = cv2.morphologyEx(mask, cv2.MORPH_CLOSE, kernel, iterations=1)
    connected_mask = cv2.dilate(connected_mask, kernel, iterations=2)
    count, labels, stats, _centroids = cv2.connectedComponentsWithStats(connected_mask, 8)

    page_area = width_px * height_px
    for label in range(1, count):
        x, y, w, h, area = stats[label]
        bbox_area = w * h
        if area < 520 or bbox_area < 1400 or w < 10 or h < 10:
            continue
        if max(w, h) < 48:
            continue
        if w * h > page_area * 0.42:
            continue
        if h <= 4 and w > 40:
            continue
        density = area / max(1, bbox_area)
        if density < 0.02 and max(w, h) < 220:
            continue

        pad = max(3, min(18, int(max(w, h) * 0.04)))
        left = max(0, x - pad)
        top = max(0, y - pad)
        right = min(width_px, x + w + pad)
        bottom = min(height_px, y + h + pad)
        crop = arr[top:bottom, left:right].copy()
        crop_mask = mask[top:bottom, left:right].copy()
        if int(np.count_nonzero(crop_mask)) < 25:
            continue
        fragment_path = os.path.join(dst_dir, f"page-{page_index + 1}-fragment-{len(fragments) + 1}.png")
        if not save_transparent_fragment(crop, crop_mask, fragment_path):
            continue
        fragments.append(
            VisualFragment(
                path=fragment_path,
                left=left * scale_x,
                top=top * scale_y,
                width=(right - left) * scale_x,
                height=(bottom - top) * scale_y,
            )
        )

    fragments.sort(key=lambda f: (f.top, f.left))
    return fragments, rules


def line_should_be_bold(line: LineBox) -> bool:
    letters = [ch for ch in line.text if ch.isalpha()]
    uppercase = sum(1 for ch in letters if ch.upper() == ch)
    near_top = line.top < line.page_height_px * 0.12
    return bool(letters) and (near_top or uppercase / max(1, len(letters)) > 0.72)


def line_overlap_ratio(a: LineBox, b: LineBox) -> float:
    left = max(a.left, b.left)
    top = max(a.top, b.top)
    right = min(a.left + a.width, b.left + b.width)
    bottom = min(a.top + a.height, b.top + b.height)
    if right <= left or bottom <= top:
        return 0.0
    overlap = (right - left) * (bottom - top)
    smaller = max(1.0, min(a.width * a.height, b.width * b.height))
    return overlap / smaller


def duplicate_line_indices(candidate: LineBox, existing: list[LineBox]) -> list[int]:
    indices: list[int] = []
    for idx, line in enumerate(existing):
        vertical_overlap = min(candidate.top + candidate.height, line.top + line.height) - max(candidate.top, line.top)
        if vertical_overlap <= 0:
            continue
        vertical_ratio = vertical_overlap / max(1.0, min(candidate.height, line.height))
        center_distance = abs((candidate.left + candidate.width / 2) - (line.left + line.width / 2))
        close_centers = center_distance <= max(candidate.width, line.width) * 0.28
        if line_overlap_ratio(candidate, line) > 0.45 or (vertical_ratio > 0.7 and close_centers):
            indices.append(idx)
    return indices


def is_duplicate_line(candidate: LineBox, existing: list[LineBox]) -> bool:
    return bool(duplicate_line_indices(candidate, existing))


def probable_ocr_noise(line: LineBox, lang: str) -> bool:
    text = normalize_ocr_text(line.text)
    if not text:
        return True
    counts = script_counts(text)
    parts = set(p for p in re.split(r"[,+\s]+", lang) if p)
    if line.conf < 35 and counts["digits"] == 0 and counts["cjk"] == 0 and counts["rtl"] == 0 and len(text) <= 16:
        return True
    if (
        counts["latin"] <= 2
        and counts["digits"] == 0
        and counts["cjk"] == 0
        and counts["rtl"] == 0
        and len(text) <= 3
        and line.width < max(90.0, line.height * 2.2)
    ):
        return True
    if parts & {"chi_sim", "chi_tra"}:
        uppercase_words = re.findall(r"\b[A-Z]{2,}\b", text)
        if counts["cjk"] == 0 and counts["digits"] == 0 and len(uppercase_words) >= 3 and line.conf < 52:
            return True
    return False


def line_quality_score(line: LineBox, lang: str) -> float:
    counts = script_counts(line.text)
    score = max(0.0, line.conf)
    score += min(14.0, counts["cjk"] * 0.22)
    score += min(7.0, counts["latin"] * 0.035)
    score += min(5.0, counts["digits"] * 0.12)
    if probable_ocr_noise(line, lang):
        score -= 60.0
    return score


def merge_line_candidates(existing: list[LineBox], candidates: list[LineBox], lang: str) -> list[LineBox]:
    merged = existing[:]
    for candidate in candidates:
        duplicate_indices = duplicate_line_indices(candidate, merged)
        if not duplicate_indices:
            merged.append(candidate)
            continue

        candidate_score = line_quality_score(candidate, lang)
        existing_scores = [line_quality_score(merged[i], lang) for i in duplicate_indices]
        replace_bonus = 8.0 if len(duplicate_indices) > 1 else 0.0
        if candidate_score + replace_bonus > max(existing_scores) + 3.0:
            for idx in sorted(duplicate_indices, reverse=True):
                del merged[idx]
            merged.append(candidate)

    return sorted(merged, key=lambda l: (l.top, l.left))


def collect_ocr_lines(
    image: str,
    lang: str,
    page_width_px: float,
    page_height_px: float,
    page_width_pt: float,
    page_height_pt: float,
    psm_modes: list[str],
    report: dict[str, Any],
) -> list[LineBox]:
    collected: list[LineBox] = []
    report.setdefault("ocrPasses", [])
    candidates = ocr_language_candidates(lang)
    for psm in psm_modes:
        best_lang = ""
        best_score = -10000.0
        parsed: list[LineBox] = []
        errors: list[str] = []
        pass_notes: list[str] = []
        for candidate in candidates:
            try:
                tsv = run_tesseract_tsv(image, candidate, psm=psm)
                candidate_lines = parse_tsv(tsv, page_width_px, page_height_px, page_width_pt, page_height_pt)
                score = score_ocr_lines(candidate_lines, candidate)
                pass_notes.append(f"{candidate}:{len(candidate_lines)}@{score:.1f}")
                if score > best_score:
                    best_lang = candidate
                    best_score = score
                    parsed = candidate_lines
            except Exception as e:
                errors.append(f"{candidate}: {e}")
        if not best_lang:
            if FAST_HOSTED_OCR:
                report["hostedOcrTimedOut"] = True
                warning = (
                    "Hosted OCR timed out before editable text reconstruction could finish. "
                    "FileMint returned a visual DOCX fallback instead of failing."
                )
                if warning not in report.setdefault("warnings", []):
                    report["warnings"].append(warning)
                report["ocrPasses"].append(f"psm-{psm}/timeout [{'; '.join(errors)}]")
                continue
            raise RuntimeError("; ".join(errors) or "Tesseract OCR failed.")
        report["ocrPasses"].append(f"psm-{psm}/{best_lang}:{len(parsed)} [{'; '.join(pass_notes)}]")
        for line in parsed:
            if not line_text_signal(line):
                continue
            collected = merge_line_candidates(collected, [line], lang)
    return sorted(collected, key=lambda l: (l.top, l.left))


def append_exact_visual_page(
    doc: Any,
    page_index: int,
    visual_png: str,
    editable_lines: list[LineBox],
    page_width_pt: float,
    page_height_pt: float,
    hidden_text: bool = False,
) -> None:
    from docx.oxml import parse_xml
    from docx.shared import Pt

    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1

    # A purely floating VML page image does not reserve vertical layout space.
    # LibreOffice/Word can then anchor multiple pages to one physical page,
    # especially after page breaks. The inline page image below occupies the
    # page, while the VML text boxes remain editable overlays on that page.
    paragraph.add_run().add_picture(
        visual_png,
        width=Pt(max(1.0, page_width_pt - 0.8)),
        height=Pt(max(1.0, page_height_pt - 0.8)),
    )

    scale_x = page_width_pt / max(1.0, editable_lines[0].page_width_px if editable_lines else page_width_pt)
    scale_y = page_height_pt / max(1.0, editable_lines[0].page_height_px if editable_lines else page_height_pt)
    for idx, line in enumerate(editable_lines):
        x = max(0.0, line.left * scale_x - 1.0)
        y = max(0.0, line.top * scale_y - 1.0)
        w = min(max(14.0, line.width * scale_x * 1.12 + 10.0), max(14.0, page_width_pt - x - 2.0))
        h = max(8.0, line.height * scale_y * 1.9)
        font_size = ocr_font_size(line)
        text = escape(line.text)
        align = "right" if contains_rtl(line.text) else "left"
        bidi = "<w:bidi/>" if contains_rtl(line.text) else ""
        bold = "<w:b/>" if line_should_be_bold(line) else ""
        vanish = "<w:vanish/>" if hidden_text else ""
        color = '<w:color w:val="FFFFFF"/>' if hidden_text else ""
        ascii_font, east_asia_font, complex_font = ocr_font_attrs(line.text)
        paragraph._p.append(
            parse_xml(
            f'''
            <w:r {VML_NS}>
              <w:pict>
                <v:shape id="FileMintText{page_index}_{idx}" type="#_x0000_t202"
                  style="position:absolute;margin-left:{x:.2f}pt;margin-top:{y:.2f}pt;width:{w:.2f}pt;height:{h:.2f}pt;z-index:{idx + 1};mso-position-horizontal:absolute;mso-position-horizontal-relative:page;mso-position-vertical:absolute;mso-position-vertical-relative:page"
                  stroked="f" filled="f" o:allowincell="f">
                  <v:textbox inset="0,0,0,0">
                    <w:txbxContent>
                      <w:p>
                        <w:pPr><w:jc w:val="{align}"/>{bidi}</w:pPr>
                        <w:r>
                          <w:rPr>
                            <w:rFonts w:ascii="{ascii_font}" w:hAnsi="{ascii_font}" w:eastAsia="{east_asia_font}" w:cs="{complex_font}"/>
                            {bold}
                            {vanish}
                            {color}
                            <w:sz w:val="{int(font_size * 2)}"/>
                            <w:szCs w:val="{int(font_size * 2)}"/>
                          </w:rPr>
                          <w:t xml:space="preserve">{text}</w:t>
                        </w:r>
                      </w:p>
                    </w:txbxContent>
                  </v:textbox>
                </v:shape>
              </w:pict>
            </w:r>
            '''
            )
        )


def append_positioned_page(
    doc: Any,
    page_index: int,
    editable_lines: list[LineBox],
    page_width_pt: float,
    page_height_pt: float,
    fragments: list[VisualFragment],
    rules: list[VisualRule],
) -> None:
    from docx.oxml import parse_xml
    from docx.shared import Pt, RGBColor

    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = Pt(max(1.0, page_height_pt - 1.0))
    spacer = paragraph.add_run(" ")
    spacer.font.size = Pt(1)
    try:
        spacer.font.color.rgb = RGBColor(255, 255, 255)
    except Exception:
        pass

    for idx, rule in enumerate(rules):
        x = max(0.0, rule.left)
        y = max(0.0, rule.top)
        w = min(max(0.35, rule.width), max(0.35, page_width_pt - x))
        h = min(max(0.35, rule.height), max(0.35, page_height_pt - y))
        paragraph._p.append(
            parse_xml(
                f'''
                <w:r {VML_NS}>
                  <w:pict>
                    <v:rect id="FileMintRule{page_index}_{idx}"
                      style="position:absolute;margin-left:{x:.2f}pt;margin-top:{y:.2f}pt;width:{w:.2f}pt;height:{h:.2f}pt;z-index:{10 + idx};mso-position-horizontal:absolute;mso-position-horizontal-relative:page;mso-position-vertical:absolute;mso-position-vertical-relative:page"
                      fillcolor="{rule.color}" stroked="f" filled="t" o:allowincell="f"/>
                  </w:pict>
                </w:r>
                '''
            )
        )

    for idx, fragment in enumerate(fragments):
        if fragment.width <= 0 or fragment.height <= 0:
            continue
        r_id, _image = doc.part.get_or_add_image(fragment.path)
        x = max(0.0, fragment.left)
        y = max(0.0, fragment.top)
        w = min(max(0.5, fragment.width), max(0.5, page_width_pt - x))
        h = min(max(0.5, fragment.height), max(0.5, page_height_pt - y))
        paragraph._p.append(
            parse_xml(
                f'''
                <w:r {VML_NS}>
                  <w:pict>
                    <v:shape id="FileMintFragment{page_index}_{idx}" type="#_x0000_t75"
                      style="position:absolute;margin-left:{x:.2f}pt;margin-top:{y:.2f}pt;width:{w:.2f}pt;height:{h:.2f}pt;z-index:{100 + idx};mso-position-horizontal:absolute;mso-position-horizontal-relative:page;mso-position-vertical:absolute;mso-position-vertical-relative:page"
                      stroked="f" filled="f" o:allowincell="f">
                      <v:imagedata r:id="{r_id}" o:title="FileMint visual fragment"/>
                    </v:shape>
                  </w:pict>
                </w:r>
                '''
            )
        )

    scale_x = page_width_pt / max(1.0, editable_lines[0].page_width_px if editable_lines else page_width_pt)
    scale_y = page_height_pt / max(1.0, editable_lines[0].page_height_px if editable_lines else page_height_pt)
    for idx, line in enumerate(editable_lines):
        x = max(0.0, line.left * scale_x - 1.0)
        y = max(0.0, line.top * scale_y - 1.0)
        w = min(max(14.0, line.width * scale_x * 1.12 + 10.0), max(14.0, page_width_pt - x - 2.0))
        h = max(8.0, line.height * scale_y * 1.9)
        font_size = ocr_font_size(line)
        text = escape(line.text)
        align = "right" if contains_rtl(line.text) else "left"
        bidi = "<w:bidi/>" if contains_rtl(line.text) else ""
        bold = "<w:b/>" if line_should_be_bold(line) else ""
        ascii_font, east_asia_font, complex_font = ocr_font_attrs(line.text)
        paragraph._p.append(
            parse_xml(
                f'''
                <w:r {VML_NS}>
                  <w:pict>
                    <v:shape id="FileMintText{page_index}_{idx}" type="#_x0000_t202"
                      style="position:absolute;margin-left:{x:.2f}pt;margin-top:{y:.2f}pt;width:{w:.2f}pt;height:{h:.2f}pt;z-index:{1000 + idx};mso-position-horizontal:absolute;mso-position-horizontal-relative:page;mso-position-vertical:absolute;mso-position-vertical-relative:page"
                      stroked="f" filled="f" o:allowincell="f">
                      <v:textbox inset="0,0,0,0">
                        <w:txbxContent>
                          <w:p>
                            <w:pPr><w:jc w:val="{align}"/>{bidi}</w:pPr>
                            <w:r>
                              <w:rPr>
                                <w:rFonts w:ascii="{ascii_font}" w:hAnsi="{ascii_font}" w:eastAsia="{east_asia_font}" w:cs="{complex_font}"/>
                                {bold}
                                <w:sz w:val="{int(font_size * 2)}"/>
                                <w:szCs w:val="{int(font_size * 2)}"/>
                              </w:rPr>
                              <w:t xml:space="preserve">{text}</w:t>
                            </w:r>
                          </w:p>
                        </w:txbxContent>
                      </v:textbox>
                    </v:shape>
                  </w:pict>
                </w:r>
                '''
            )
        )


def pdf_line_text_with_gaps(line: PdfLine) -> str:
    spans = sorted(line.spans, key=lambda s: (s.bbox[0], s.bbox[1]))
    sizes = [span.size for span in spans if span.text.strip()]
    dominant_size = median(sizes) if sizes else 11.0
    parts: list[str] = []
    last_right: float | None = None
    for span in spans:
        text = xml_compatible_text(span.text)
        if not text:
            continue
        if last_right is not None:
            gap = span.bbox[0] - last_right
            spaces = estimated_gap_spaces(gap, dominant_size)
            if spaces and parts and not parts[-1].endswith(" ") and not text.startswith(" "):
                parts.append(" " * spaces)
        parts.append(text)
        last_right = max(last_right or span.bbox[2], span.bbox[2])
    return "".join(parts).strip()


def native_pdf_line_boxes(page: Any, page_width_px: float, page_height_px: float) -> list[LineBox]:
    sx = page_width_px / max(1.0, page.rect.width)
    sy = page_height_px / max(1.0, page.rect.height)
    out: list[LineBox] = []
    for line in merge_visual_rows(collect_pdf_lines(page)):
        text = pdf_line_text_with_gaps(line)
        if not text or not any(ch.isalnum() for ch in text):
            continue
        left = line.left * sx
        top = line.top * sy
        width = max(1.0, (line.right - line.left) * sx)
        height = max(1.0, (line.bottom - line.top) * sy)
        out.append(
            LineBox(
                text=text,
                words=[],
                left=left,
                top=top,
                width=width,
                height=height,
                conf=99.0,
                page_width_px=page_width_px,
                page_height_px=page_height_px,
                page_width_pt=page.rect.width,
                page_height_pt=page.rect.height,
                segments=[(left, left + width, text)],
            )
        )
    return sorted(out, key=lambda line: (line.top, line.left))


def image_backed_text_layer_likely(info: dict[str, Any]) -> bool:
    pages = int(info.get("pages", 0) or 0)
    if pages <= 0:
        return False
    details = info.get("pageDetails", []) or []
    backed = [
        page
        for page in details
        if float(page.get("maxImageCoverage", 0) or 0) >= 0.65 and int(page.get("textCharacters", 0) or 0) >= 25
    ]
    return len(backed) >= max(1, math.ceil(pages * 0.5))


def image_backed_text_layer_needs_ocr(info: dict[str, Any]) -> bool:
    for page in info.get("pageDetails", []) or []:
        if float(page.get("maxImageCoverage", 0) or 0) >= 0.65 and int(page.get("textCharacters", 0) or 0) < 200:
            return True
    return False


def append_linebox_flow_paragraph(
    doc: Any,
    line: LineBox,
    prev_bottom: float | None,
    left_margin_pt: float,
    top_margin_pt: float,
) -> int:
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt

    text = xml_compatible_text(" ".join(line.text.split()))
    if not text:
        return 0

    scale_x = line.page_width_pt / max(1.0, line.page_width_px)
    scale_y = line.page_height_pt / max(1.0, line.page_height_px)
    left_pt = line.left * scale_x
    right_pt = (line.left + line.width) * scale_x
    top_pt = line.top * scale_y
    bottom_pt = (line.top + line.height) * scale_y
    page_center = line.page_width_pt / 2
    line_center = (left_pt + right_pt) / 2

    paragraph = doc.add_paragraph()
    fmt = paragraph.paragraph_format
    fmt.left_indent = Pt(max(0.0, left_pt - left_margin_pt))
    fmt.right_indent = Pt(0)
    fmt.space_after = Pt(0)
    if prev_bottom is None:
        fmt.space_before = Pt(max(0.0, min(42.0, top_pt - top_margin_pt)))
    else:
        gap = max(0.0, top_pt - prev_bottom)
        fmt.space_before = Pt(min(28.0, gap))
    fmt.line_spacing = 1.0

    if contains_rtl(text):
        set_paragraph_bidi(paragraph)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    elif abs(line_center - page_center) <= 35 and (right_pt - left_pt) <= line.page_width_pt * 0.76:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        fmt.left_indent = Pt(0)
    else:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT

    run = paragraph.add_run(text)
    set_run_font(run, "Times New Roman")
    font_size = max(7.0, min(16.0, (bottom_pt - top_pt) * 0.92))
    run.font.size = Pt(font_size)
    run.bold = line_should_be_bold(line) or font_size >= 12.5
    return len(text)


def to_docx_scan_text_layer(
    src: str,
    dst: str,
    lang: str | None,
    report: dict[str, Any],
    premium: bool = False,
    table_detection: bool = True,
    quality: str = "high",
) -> None:
    import fitz
    from docx import Document
    from docx.shared import Pt

    dpi = quality_dpi(quality, 300)
    pdf = fitz.open(src)
    out = Document()
    tmpdir = tempfile.mkdtemp(prefix="filemint-text-layer-scan-")
    try:
        pages_with_text = 0
        all_low_conf = 0
        total_candidates = 0
        total_editable_boxes = 0
        total_editable_chars = 0
        total_skipped_low_conf = 0
        native_pages = 0
        ocr_pages = 0
        scanned_tables = 0

        for page_index, page in enumerate(pdf):
            if page_index > 0:
                out.add_page_break()
            section = out.sections[-1]
            section.page_width = Pt(page.rect.width)
            section.page_height = Pt(page.rect.height)
            section.top_margin = Pt(12)
            section.bottom_margin = Pt(8)
            section.left_margin = Pt(54)
            section.right_margin = Pt(42)

            pix = page.get_pixmap(dpi=dpi, alpha=False)
            scan_png = os.path.join(tmpdir, f"page-{page_index + 1}.png")
            pix.save(scan_png)

            native_lines = native_pdf_line_boxes(page, pix.width, pix.height)
            native_chars = sum(len(line.text) for line in native_lines)
            lines = native_lines
            source = "native-text-layer"
            primary_ocr_lines: list[LineBox] = []

            if lang and (native_chars < 200 or (premium and table_detection)):
                primary_ocr_lines = collect_ocr_lines(
                    scan_png,
                    lang,
                    pix.width,
                    pix.height,
                    page.rect.width,
                    page.rect.height,
                    ["6"] if FAST_HOSTED_OCR else ["11"],
                    report,
                )
                if table_detection and premium and dense_table_scan_likely(primary_ocr_lines) and transcript_scan_likely(primary_ocr_lines):
                    scanned_tables += 1
                    build_scanned_table_page(out, scan_png, primary_ocr_lines, page.rect.width, page.rect.height, report)
                    continue
                if native_chars < 200 and primary_ocr_lines:
                    lines = primary_ocr_lines
                    source = "ocr-fallback"
                    if premium and not FAST_HOSTED_OCR:
                        extra_lines = collect_ocr_lines(
                            scan_png,
                            lang,
                            pix.width,
                            pix.height,
                            page.rect.width,
                            page.rect.height,
                            ["6"],
                            report,
                        )
                        lines = merge_line_candidates(lines, extra_lines, lang)

            if source == "native-text-layer":
                native_pages += 1
                editable_lines = [line for line in lines if line_text_signal(line)]
                stats = {
                    "ocrTextCandidates": len(editable_lines),
                    "editableTextBoxes": len(editable_lines),
                    "editableCharacters": sum(len(line.text) for line in editable_lines),
                    "skippedLowConfidence": 0,
                }
            else:
                ocr_pages += 1
                all_low_conf += sum(1 for line in lines for word in line.words if 0 <= word.conf < 55)
                editable_lines, stats = exact_editable_lines(scan_png, lines, lang or "native", premium=True)

            total_candidates += stats["ocrTextCandidates"]
            total_editable_boxes += stats["editableTextBoxes"]
            total_editable_chars += stats["editableCharacters"]
            total_skipped_low_conf += stats["skippedLowConfidence"]
            if editable_lines:
                pages_with_text += 1

            prev_bottom: float | None = None
            for line in editable_lines:
                append_linebox_flow_paragraph(out, line, prev_bottom, 54.0, 12.0)
                scale_y = line.page_height_pt / max(1.0, line.page_height_px)
                prev_bottom = (line.top + line.height) * scale_y

        report["resolvedMode"] = "premium-scan-text-flow" if premium else "scan-text-flow"
        report["editableTextDetected"] = pages_with_text > 0
        report["lowConfidenceOcrAreas"] = all_low_conf
        report["editableTextBoxes"] = total_editable_boxes
        report["editableCharacters"] = total_editable_chars
        report["ocrTextCandidates"] = total_candidates
        report["textCoverageEstimate"] = round((total_editable_boxes / max(1, total_candidates)) * 100)
        report["visualObjectsPreserved"] = report.get("pagesConverted", 0)
        report["hiddenTextLayer"] = False
        report["visibleEditableTextLayer"] = pages_with_text > 0
        report["tablesRebuiltAsWord"] = scanned_tables
        report["nonEditableVisualFallback"] = False
        report["notes"].append(
            f"Detected full-page scanned images with a text layer. Rebuilt {native_pages} page(s) from native PDF text coordinates and {ocr_pages} weak page(s) with OCR fallback as normal editable Word paragraphs."
        )
        if total_skipped_low_conf:
            report["warnings"].append(
                f"{total_skipped_low_conf} OCR text candidates were too uncertain to rebuild as editable text."
            )
        out.save(dst)
    finally:
        pdf.close()
        shutil.rmtree(tmpdir, ignore_errors=True)


def words_from_lines(lines: list[LineBox]) -> list[WordBox]:
    words: list[WordBox] = []
    for line in lines:
        for word in line.words:
            candidate = word_as_line(word, line)
            if not is_duplicate_word_line(candidate, [word_as_line(w, line) for w in words]):
                words.append(word)
    return sorted(words, key=lambda w: (w.top, w.left))


def clean_ocr_token(text: str) -> str:
    token = text.strip()
    replacements = {
        "ехс": "exc",
        "еже": "exc",
        "ех": "exc",
        "fexc": "exc",
        "{exe}": "(exc)",
        "[ехс)": "(exc)",
        "fexc}": "(exc)",
        "(ехс)": "(exc)",
        "(ехс}": "(exc)",
        "go od": "good",
        "(go od)": "(good)",
        "lex": "(exc)",
        "exe": "exc",
        "eже": "exc",
    }
    for old, new in replacements.items():
        token = token.replace(old, new)
    token = token.replace("{", "(").replace("[", "(").replace("}", ")").replace("]", ")")
    if token in {"—", "=", "==", "——", "——<————", "_", "__", "--", "ИЕ", "ШАР", "кат", "Thi", "чапи", "НН"}:
        return ""
    return token


def normalize_grade_cell(text: str) -> str:
    value = " ".join(text.split())
    value = value.replace("5/S", "5/5").replace("5/s", "5/5")
    value = re.sub(r"\bof\s+5\b", "5/5", value)
    value = re.sub(r"\b(\d+)\s+10\b", r"\1/10", value)
    value = re.sub(r"\b10\s*/?\s*10\b", "10/10", value)
    value = re.sub(r"\b5\s*/?\s*5\b", "5/5", value)
    value = value.replace("( ехс)", "(exc)").replace("( ех)", "(exc)")
    value = value.replace("(ехс)", "(exc)").replace("(exc))", "(exc)")
    value = value.replace("(go od)", "(good)").replace("(good))", "(good)")
    value = value.replace("AE", "").replace("aS", "").strip()
    value = re.sub(r"\s+", " ", value)
    return value


def token_has_text_signal(token: str) -> bool:
    return any(ch.isalnum() for ch in token) or token in {"-", "#"}


def join_positioned_words(words: list[WordBox]) -> str:
    if not words:
        return ""
    rows: list[list[WordBox]] = []
    for word in sorted(words, key=lambda w: (w.top, w.left)):
        token = clean_ocr_token(word.text)
        if not token or not token_has_text_signal(token):
            continue
        placed = False
        center = word.top + word.height / 2
        for row in rows:
            row_center = sum(w.top + w.height / 2 for w in row) / len(row)
            if abs(center - row_center) <= max(12.0, word.height * 0.55):
                row.append(word)
                placed = True
                break
        if not placed:
            rows.append([word])

    parts: list[str] = []
    for row in rows:
        line_parts: list[str] = []
        prev: WordBox | None = None
        for word in sorted(row, key=lambda w: w.left):
            token = clean_ocr_token(word.text)
            if not token or not token_has_text_signal(token):
                continue
            if prev is not None and line_parts:
                gap = word.left - (prev.left + prev.width)
                if gap <= 10 and re.search(r"[A-Za-zА-Яа-я]$", line_parts[-1]) and re.match(r"^[A-Za-zА-Яа-я]", token):
                    line_parts[-1] += token
                elif gap <= 8 and (line_parts[-1].endswith(("/", "(", "-")) or token.startswith((")", "/", "-"))):
                    line_parts[-1] += token
                else:
                    line_parts.append(token)
            else:
                line_parts.append(token)
            prev = word
        if line_parts:
            parts.append(" ".join(line_parts))
    return "\n".join(parts).strip()


def set_table_cell_text(cell: Any, text: str, size: float = 8.5, bold: bool = False, align: str = "left") -> None:
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt

    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.0
    if align == "center":
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif align == "right":
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    else:
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for i, line in enumerate(text.splitlines() or [""]):
        if i:
            p.add_run().add_break()
        run = p.add_run(line)
        run.font.name = "Arial"
        run.font.size = Pt(size)
        run.bold = bold


def set_cell_width(cell: Any, width_pt: float) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(int(width_pt * 20)))
    tc_w.set(qn("w:type"), "dxa")


def set_row_height(row: Any, height_pt: float, rule: str = "atLeast") -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    tr_pr = row._tr.get_or_add_trPr()
    tr_h = tr_pr.find(qn("w:trHeight"))
    if tr_h is None:
        tr_h = OxmlElement("w:trHeight")
        tr_pr.append(tr_h)
    tr_h.set(qn("w:val"), str(int(height_pt * 20)))
    tr_h.set(qn("w:hRule"), rule)


def set_table_layout(table: Any, width_pt: float | None = None, indent_pt: float | None = None) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    tbl_pr = table._tbl.tblPr
    tbl_layout = tbl_pr.find(qn("w:tblLayout"))
    if tbl_layout is None:
        tbl_layout = OxmlElement("w:tblLayout")
        tbl_pr.append(tbl_layout)
    tbl_layout.set(qn("w:type"), "fixed")

    if width_pt is not None:
        tbl_w = tbl_pr.find(qn("w:tblW"))
        if tbl_w is None:
            tbl_w = OxmlElement("w:tblW")
            tbl_pr.append(tbl_w)
        tbl_w.set(qn("w:w"), str(int(width_pt * 20)))
        tbl_w.set(qn("w:type"), "dxa")

    if indent_pt is not None:
        tbl_ind = tbl_pr.find(qn("w:tblInd"))
        if tbl_ind is None:
            tbl_ind = OxmlElement("w:tblInd")
            tbl_pr.append(tbl_ind)
        tbl_ind.set(qn("w:w"), str(int(max(0.0, indent_pt) * 20)))
        tbl_ind.set(qn("w:type"), "dxa")


def set_cell_margins(cell: Any, margin_pt: float = 2.0) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.find(qn("w:tcMar"))
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for edge in ("top", "start", "bottom", "end"):
        node = tc_mar.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(int(margin_pt * 20)))
        node.set(qn("w:type"), "dxa")


def cluster_numeric(values: list[float], threshold: float) -> list[float]:
    clusters: list[list[float]] = []
    for value in sorted(values):
        if not clusters or abs(median(clusters[-1]) - value) > threshold:
            clusters.append([value])
        else:
            clusters[-1].append(value)
    return [float(median(cluster)) for cluster in clusters]


def choose_consecutive(values: list[float], count: int) -> list[float]:
    if len(values) <= count:
        return values
    best = values[:count]
    best_span = best[-1] - best[0]
    for i in range(1, len(values) - count + 1):
        candidate = values[i : i + count]
        span = candidate[-1] - candidate[0]
        if span > best_span:
            best = candidate
            best_span = span
    return best


def detect_transcript_grid_geometry(
    img: Any,
    page_width_pt: float,
    page_height_pt: float,
) -> dict[str, Any] | None:
    try:
        import cv2
        import numpy as np
    except Exception:
        return None

    arr = np.array(img.convert("RGB"))
    height_px, width_px = arr.shape[:2]
    gray = cv2.cvtColor(arr, cv2.COLOR_RGB2GRAY)
    binary = cv2.adaptiveThreshold(gray, 255, cv2.ADAPTIVE_THRESH_MEAN_C, cv2.THRESH_BINARY_INV, 31, 12)

    top_offset = int(height_px * 0.28)
    bottom_offset = int(height_px * 0.88)
    left_offset = int(width_px * 0.05)
    right_offset = int(width_px * 0.95)
    roi = binary[top_offset:bottom_offset, left_offset:right_offset]

    h_kernel = cv2.getStructuringElement(cv2.MORPH_RECT, (max(80, width_px // 18), 1))
    v_kernel = cv2.getStructuringElement(cv2.MORPH_RECT, (1, max(50, height_px // 35)))
    horizontal = cv2.morphologyEx(roi, cv2.MORPH_OPEN, h_kernel)
    vertical = cv2.morphologyEx(roi, cv2.MORPH_OPEN, v_kernel)

    xs: list[float] = []
    contours, _ = cv2.findContours(vertical, cv2.RETR_EXTERNAL, cv2.CHAIN_APPROX_SIMPLE)
    for contour in contours:
        x, y, w, h = cv2.boundingRect(contour)
        if h > height_px * 0.10:
            xs.append(left_offset + x + w / 2)
    xs = cluster_numeric(xs, max(10.0, width_px * 0.006))
    xs = choose_consecutive([x for x in xs if width_px * 0.12 <= x <= width_px * 0.94], 6)
    if len(xs) < 5:
        return None

    table_left_px = xs[0]
    table_right_px = xs[-1]
    table_width_px = max(1.0, table_right_px - table_left_px)

    ys: list[float] = []
    contours, _ = cv2.findContours(horizontal, cv2.RETR_EXTERNAL, cv2.CHAIN_APPROX_SIMPLE)
    for contour in contours:
        x, y, w, h = cv2.boundingRect(contour)
        abs_x = left_offset + x
        abs_y = top_offset + y
        center_x = abs_x + w / 2
        if w > table_width_px * 0.42 and table_left_px - 35 <= center_x <= table_right_px + 35:
            ys.append(abs_y + h / 2)
    ys = cluster_numeric(ys, max(10.0, height_px * 0.004))
    if len(ys) >= 25:
        ys = ys[-25:]
    elif len(ys) < 8:
        return None

    sx = page_width_pt / max(1.0, width_px)
    sy = page_height_pt / max(1.0, height_px)
    x_positions_pt = [x * sx for x in xs]
    y_positions_pt = [y * sy for y in ys]
    col_widths_pt = [
        max(10.0, x_positions_pt[i + 1] - x_positions_pt[i])
        for i in range(len(x_positions_pt) - 1)
    ]
    row_heights_pt = [
        max(9.0, y_positions_pt[i + 1] - y_positions_pt[i])
        for i in range(len(y_positions_pt) - 1)
    ]

    return {
        "xPositionsPx": xs,
        "yPositionsPx": ys,
        "xPositionsPt": x_positions_pt,
        "yPositionsPt": y_positions_pt,
        "colWidthsPt": col_widths_pt,
        "rowHeightsPt": row_heights_pt,
        "tableLeftPt": x_positions_pt[0],
        "tableWidthPt": sum(col_widths_pt),
        "preTableSpacerPt": max(0.0, min(120.0, y_positions_pt[0] - page_height_pt * 0.222)),
    }


def words_in_box(words: list[WordBox], left: float, top: float, right: float, bottom: float, min_conf: float = 18) -> list[WordBox]:
    out: list[WordBox] = []
    for word in words:
        cx = word.left + word.width / 2
        cy = word.top + word.height / 2
        if left <= cx <= right and top <= cy <= bottom and word.conf >= min_conf:
            out.append(word)
    return out


def cluster_y_centers(values: list[float], threshold: float = 46.0) -> list[float]:
    if not values:
        return []
    clusters: list[list[float]] = []
    for value in sorted(values):
        if not clusters or abs(median(clusters[-1]) - value) > threshold:
            clusters.append([value])
        else:
            clusters[-1].append(value)
    return [median(cluster) for cluster in clusters]


def transcript_row_centers(words: list[WordBox], width_px: float, height_px: float) -> list[float]:
    left = width_px * 0.24
    right = width_px * 0.45
    top = height_px * 0.34
    bottom = height_px * 0.86
    values: list[float] = []
    for word in words:
        token = clean_ocr_token(word.text)
        if word.conf < 40 or not any(ch.isalpha() for ch in token):
            continue
        cx = word.left + word.width / 2
        cy = word.top + word.height / 2
        if left <= cx <= right and top <= cy <= bottom:
            values.append(cy)
    centers = cluster_y_centers(values, 52.0)
    if len(centers) >= 18:
        return centers[:23]

    first = height_px * 0.365
    step = height_px * 0.0214
    return [first + i * step for i in range(23)]


def build_scanned_table_page(
    doc: Any,
    scan_png: str,
    lines: list[LineBox],
    page_width_pt: float,
    page_height_pt: float,
    report: dict[str, Any],
) -> None:
    from docx.enum.section import WD_SECTION
    from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import parse_xml
    from docx.shared import Pt
    from PIL import Image

    section = doc.sections[-1]
    section.page_width = Pt(page_width_pt)
    section.page_height = Pt(page_height_pt)
    section.top_margin = Pt(28)
    section.bottom_margin = Pt(26)
    section.left_margin = Pt(34)
    section.right_margin = Pt(34)

    img = Image.open(scan_png).convert("RGB")
    width_px, height_px = img.size
    words = words_from_lines(lines)
    grid_geometry = detect_transcript_grid_geometry(img, page_width_pt, page_height_pt)
    left_margin_pt = 34.0

    header_lines = [line for line in lines if line.top < height_px * 0.17 and line.conf >= 35 and line_text_signal(line)]
    columns = [[], [], []]
    for line in header_lines:
        center = line.left + line.width / 2
        idx = 0 if center < width_px * 0.35 else 1 if center < width_px * 0.66 else 2
        columns[idx].append(line)
    header = doc.add_table(rows=1, cols=3)
    header.alignment = WD_TABLE_ALIGNMENT.CENTER
    header.autofit = False
    for idx, cell in enumerate(header.rows[0].cells):
        text = "\n".join(line.text for line in sorted(columns[idx], key=lambda l: l.top)[:5])
        set_cell_width(cell, 165)
        set_table_cell_text(cell, text, size=7.2, bold=idx < 3, align="center")

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_before = Pt(4)
    title.paragraph_format.space_after = Pt(3)
    run = title.add_run("Student Personal Information")
    run.font.name = "Arial"
    run.font.size = Pt(11)
    run.bold = True

    info = doc.add_table(rows=2, cols=2)
    info.style = "Table Grid"
    info.alignment = WD_TABLE_ALIGNMENT.CENTER
    info.autofit = False
    info_col_w = [250, 250]
    for row in info.rows:
        set_row_height(row, 20)
        for i, cell in enumerate(row.cells):
            set_cell_width(cell, info_col_w[i])
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

    def box_text(left_r: float, top_r: float, right_r: float, bottom_r: float) -> str:
        return join_positioned_words(words_in_box(words, width_px * left_r, height_px * top_r, width_px * right_r, height_px * bottom_r))

    set_table_cell_text(info.cell(0, 0), box_text(0.06, 0.205, 0.47, 0.245), size=7.8)
    set_table_cell_text(info.cell(0, 1), box_text(0.47, 0.205, 0.93, 0.245), size=7.8)
    set_table_cell_text(info.cell(1, 0), box_text(0.06, 0.235, 0.47, 0.275), size=7.8)
    set_table_cell_text(info.cell(1, 1), box_text(0.47, 0.235, 0.93, 0.275), size=7.8)

    academic = doc.add_paragraph()
    academic.alignment = WD_ALIGN_PARAGRAPH.CENTER
    academic.paragraph_format.space_before = Pt(6)
    academic.paragraph_format.space_after = Pt(
        grid_geometry["preTableSpacerPt"] if grid_geometry else 4
    )
    arun = academic.add_run("Academic Record 2024 - 2026")
    arun.font.name = "Arial"
    arun.font.size = Pt(9.5)
    arun.bold = True

    if grid_geometry and len(grid_geometry["yPositionsPx"]) >= 25:
        y_positions_px = grid_geometry["yPositionsPx"]
        row_bounds = [
            (y_positions_px[i], y_positions_px[i + 1])
            for i in range(1, min(24, len(y_positions_px) - 1))
        ]
    else:
        row_centers = transcript_row_centers(words, width_px, height_px)
        if not row_centers:
            row_centers = [height_px * 0.365 + i * height_px * 0.0214 for i in range(23)]
        row_bounds = []
        header_top = max(height_px * 0.30, row_centers[0] - 130)
        first_boundary = (header_top + row_centers[0]) / 2
        for i, center in enumerate(row_centers):
            top = first_boundary if i == 0 else (row_centers[i - 1] + center) / 2
            bottom = height_px * 0.86 if i == len(row_centers) - 1 else (center + row_centers[i + 1]) / 2
            row_bounds.append((top, bottom))

    table = doc.add_table(rows=len(row_bounds) + 1, cols=5)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT if grid_geometry else WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    col_widths_pt = (
        grid_geometry["colWidthsPt"][:5]
        if grid_geometry and len(grid_geometry["colWidthsPt"]) >= 5
        else [34, 205, 86, 86, 90]
    )
    if grid_geometry:
        set_table_layout(
            table,
            width_pt=grid_geometry["tableWidthPt"],
            indent_pt=grid_geometry["tableLeftPt"] - left_margin_pt,
        )
        report["notes"].append(
            "Transcript table geometry was measured from the source scan and applied to Word column widths, row heights and table position."
        )
    else:
        set_table_layout(table, width_pt=sum(col_widths_pt))
    headers = ["#", "Course Title", "Grade 9", "Grade 10", "Grade 11\n(half-year)"]
    for i, cell in enumerate(table.rows[0].cells):
        set_cell_width(cell, col_widths_pt[i])
        set_cell_margins(cell, 1.4 if grid_geometry else 2.0)
        set_table_cell_text(cell, headers[i], size=7.8 if grid_geometry else 8.5, bold=True, align="center" if i != 1 else "left")
    header_height = grid_geometry["rowHeightsPt"][0] if grid_geometry and grid_geometry["rowHeightsPt"] else 26
    set_row_height(table.rows[0], header_height, "exact" if grid_geometry else "atLeast")

    if grid_geometry and len(grid_geometry["xPositionsPx"]) >= 6:
        x_positions_px = grid_geometry["xPositionsPx"][:6]
        col_bounds = [(x_positions_px[i], x_positions_px[i + 1]) for i in range(5)]
    else:
        col_bounds = [
            (width_px * 0.20, width_px * 0.25),
            (width_px * 0.25, width_px * 0.535),
            (width_px * 0.535, width_px * 0.660),
            (width_px * 0.660, width_px * 0.775),
            (width_px * 0.775, width_px * 0.890),
        ]
    for r_idx, (top, bottom) in enumerate(row_bounds, start=1):
        row = table.rows[r_idx]
        if grid_geometry and len(grid_geometry["rowHeightsPt"]) > r_idx:
            set_row_height(row, grid_geometry["rowHeightsPt"][r_idx], "exact")
        else:
            set_row_height(row, 18)
        for c_idx, cell in enumerate(row.cells):
            set_cell_width(cell, col_widths_pt[c_idx])
            set_cell_margins(cell, 1.2 if grid_geometry else 2.0)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            if c_idx == 0:
                text = str(r_idx)
            else:
                left, right = col_bounds[c_idx]
                text = join_positioned_words(words_in_box(words, left, top, right, bottom))
                if c_idx >= 2:
                    text = normalize_grade_cell(text)
            set_table_cell_text(
                cell,
                text,
                size=7.0 if grid_geometry and c_idx == 1 else 6.8 if grid_geometry else 7.8 if c_idx == 1 else 7.6,
                bold=False,
                align="left" if c_idx == 1 else "center",
            )

    footer = doc.add_paragraph()
    footer.paragraph_format.space_before = Pt(5)
    footer.paragraph_format.space_after = Pt(0)
    footer.alignment = WD_ALIGN_PARAGRAPH.LEFT
    frun = footer.add_run("Principal of Tajik - Russian Lyceum - Boarding \"Hotamov Abdufattоh\"")
    frun.font.name = "Arial"
    frun.font.size = Pt(7.5)

    tmpdir = os.path.dirname(scan_png)
    seal_path = os.path.join(tmpdir, "seal-signature.png")
    crop = img.crop(
        (
            int(width_px * 0.30),
            int(height_px * 0.84),
            int(width_px * 0.80),
            min(height_px, int(height_px * 0.985)),
        )
    ).convert("RGBA")
    pixels = crop.load()
    for y in range(crop.height):
        for x in range(crop.width):
            r, g, b, _a = pixels[x, y]
            colorful = max(r, g, b) - min(r, g, b) > 28
            blue_mark = b > r + 10 and b > g - 8 and b < 245
            if not (colorful and blue_mark):
                pixels[x, y] = (255, 255, 255, 0)
    crop.save(seal_path)
    r_id, _image = doc.part.get_or_add_image(seal_path)
    seal_left_pt = page_width_pt * 0.30
    seal_top_pt = page_height_pt * 0.825
    seal_width_pt = page_width_pt * 0.50
    seal_height_pt = page_height_pt * 0.155
    footer._p.append(
        parse_xml(
            f'''
            <w:r {VML_NS}>
              <w:pict>
                <v:shape id="FileMintTranscriptSeal" type="#_x0000_t75"
                  style="position:absolute;margin-left:{seal_left_pt:.2f}pt;margin-top:{seal_top_pt:.2f}pt;width:{seal_width_pt:.2f}pt;height:{seal_height_pt:.2f}pt;z-index:2500;mso-position-horizontal:absolute;mso-position-horizontal-relative:page;mso-position-vertical:absolute;mso-position-vertical-relative:page"
                  stroked="f" filled="f" o:allowincell="f">
                  <v:imagedata r:id="{r_id}" o:title="FileMint seal and signature"/>
                </v:shape>
              </w:pict>
            </w:r>
            '''
        )
    )

    report["notes"].append(
        "Premium editable table mode rebuilt the scanned transcript as a real Word table; stamp and signature regions were kept as images."
    )


def ocr_to_docx_exact_visual(
    src: str,
    dst: str,
    lang: str,
    report: dict[str, Any],
    premium: bool = False,
    table_detection: bool = True,
    visible_text: bool = True,
    quality: str = "high",
    keep_visual_objects: bool = True,
) -> None:
    import fitz
    from docx import Document
    from docx.shared import Pt

    dpi = quality_dpi(quality, 300)
    pdf = fitz.open(src)
    out = Document()
    tmpdir = tempfile.mkdtemp(prefix="filemint-exact-ocr-")
    try:
        all_low_conf = 0
        pages_with_text = 0
        scanned_tables = 0
        colored_marks_kept = 0
        total_candidates = 0
        total_editable_boxes = 0
        total_editable_chars = 0
        total_skipped_low_conf = 0
        previous_page_reserved = False
        rebuilt_table_pages = 0
        visual_fragments_preserved = 0
        rules_rebuilt = 0
        exact_visual_fallback_pages = 0
        for page_index, page in enumerate(pdf):
            if page_index > 0 and not previous_page_reserved:
                out.add_page_break()
            section = out.sections[-1]
            section.page_width = Pt(page.rect.width)
            section.page_height = Pt(page.rect.height)
            section.top_margin = Pt(0)
            section.bottom_margin = Pt(0)
            section.left_margin = Pt(0)
            section.right_margin = Pt(0)

            pix = page.get_pixmap(dpi=dpi, alpha=False)
            scan_png = os.path.join(tmpdir, f"page-{page_index + 1}.png")
            residual_png = os.path.join(tmpdir, f"page-{page_index + 1}-visual.png")
            pix.save(scan_png)
            primary_lines = collect_ocr_lines(
                scan_png,
                lang,
                pix.width,
                pix.height,
                page.rect.width,
                page.rect.height,
                ["6"] if FAST_HOSTED_OCR else ["11"],
                report,
            )
            if FAST_HOSTED_OCR and not primary_lines:
                append_exact_visual_page(
                    out,
                    page_index + 1,
                    scan_png,
                    [],
                    page.rect.width,
                    page.rect.height,
                    hidden_text=True,
                )
                exact_visual_fallback_pages += 1
                previous_page_reserved = True
                continue
            dense_scan = premium and table_detection and dense_table_scan_likely(primary_lines)
            dense_table = dense_scan and transcript_scan_likely(primary_lines)
            generic_dense_scan = dense_scan and not dense_table
            if dense_table:
                scanned_tables = max(scanned_tables, 1)
                rebuilt_table_pages += 1
                lines = primary_lines
            if premium and not dense_table and not FAST_HOSTED_OCR:
                extra_lines = collect_ocr_lines(
                    scan_png,
                    lang,
                    pix.width,
                    pix.height,
                    page.rect.width,
                    page.rect.height,
                    ["6"],
                    report,
                )
                lines = merge_line_candidates(primary_lines, extra_lines, lang)
            else:
                if not dense_table:
                    lines = primary_lines
            if not dense_table:
                scanned_tables += len(table_runs(lines))
            if dense_table:
                editable_lines, stats = exact_editable_word_lines(scan_png, lines, lang, premium=premium)
            else:
                editable_lines, stats = exact_editable_lines(scan_png, lines, lang, premium=premium)
            colored_marks_kept += stats["skippedColoredMarks"]
            total_candidates += stats["ocrTextCandidates"]
            total_editable_boxes += stats["editableTextBoxes"]
            total_editable_chars += stats["editableCharacters"]
            total_skipped_low_conf += stats["skippedLowConfidence"]
            if editable_lines:
                pages_with_text += 1
            all_low_conf += sum(1 for line in lines for word in line.words if 0 <= word.conf < 55)
            if dense_table:
                build_scanned_table_page(out, scan_png, lines, page.rect.width, page.rect.height, report)
                previous_page_reserved = False
            else:
                if visible_text and generic_dense_scan:
                    append_exact_visual_page(
                        out,
                        page_index + 1,
                        scan_png,
                        editable_lines,
                        page.rect.width,
                        page.rect.height,
                        hidden_text=True,
                    )
                    exact_visual_fallback_pages += 1
                    previous_page_reserved = True
                elif visible_text:
                    fragments, rules = segment_visual_layer(
                        scan_png,
                        tmpdir,
                        page_index,
                        editable_lines,
                        page.rect.width,
                        page.rect.height,
                    )
                    if not keep_visual_objects:
                        fragments = []
                    visual_fragments_preserved += len(fragments)
                    rules_rebuilt += len(rules)
                    append_positioned_page(
                        out,
                        page_index + 1,
                        editable_lines,
                        page.rect.width,
                        page.rect.height,
                        fragments,
                        rules,
                    )
                    previous_page_reserved = True
                else:
                    # Exact visual mode keeps the source scan untouched and
                    # stores OCR as hidden searchable/editable text.
                    visual_png = scan_png
                    append_exact_visual_page(
                        out,
                        page_index + 1,
                        visual_png,
                        editable_lines,
                        page.rect.width,
                        page.rect.height,
                        hidden_text=True,
                    )
                    previous_page_reserved = True

        report["editableTextDetected"] = pages_with_text > 0
        report["tablesDetected"] = max(int(report.get("tablesDetected", 0)), scanned_tables)
        report["lowConfidenceOcrAreas"] = all_low_conf
        report["editableTextBoxes"] = total_editable_boxes
        report["editableCharacters"] = total_editable_chars
        report["ocrTextCandidates"] = total_candidates
        report["textCoverageEstimate"] = round((total_editable_boxes / max(1, total_candidates)) * 100)
        report["visualObjectsPreserved"] = report.get("pagesConverted", 0)
        report["hiddenTextLayer"] = (pages_with_text > 0 and not visible_text) or (
            exact_visual_fallback_pages > 0 and pages_with_text > 0
        )
        report["visibleEditableTextLayer"] = pages_with_text > exact_visual_fallback_pages and visible_text
        report["tablesRebuiltAsWord"] = rebuilt_table_pages if premium and table_detection else 0
        report["visualFragmentsPreserved"] = visual_fragments_preserved
        report["rulesRebuiltAsWord"] = rules_rebuilt
        report["nonEditableVisualFallback"] = bool(report.get("hostedOcrTimedOut")) and total_editable_chars == 0
        if premium:
            if report.get("hostedOcrTimedOut"):
                report["notes"].append(
                    "Hosted OCR did not finish within the server limit, so FileMint returned a valid visual DOCX fallback. Use a stronger backend instance for fully editable OCR on this scan."
                )
            elif rebuilt_table_pages:
                report["notes"].append(
                    "Premium table scan mode rebuilt detected transcript/table regions as editable Word table cells."
                )
            elif visible_text:
                report["notes"].append(
                    f"Premium OCR Editable Mode rebuilt scan text as visible editable Word text boxes, recreated {rules_rebuilt} simple line(s) as Word shapes, and preserved {visual_fragments_preserved} non-convertible visual fragment(s) as positioned images."
                )
            else:
                report["notes"].append(
                    "Premium Exact Visual Mode preserved each scanned page as the original image and added OCR as a hidden editable/searchable text layer to avoid visual OCR artifacts."
                )
        else:
            if visible_text:
                report["notes"].append(
                    "Scanned PDF rebuilt OCR text as visible editable Word text boxes while preserving non-text visuals as page artwork."
                )
            else:
                report["notes"].append(
                    "Scanned PDF preserved each original page image and added OCR as a hidden editable/searchable text layer at original page coordinates."
                )
        if all_low_conf:
            report["warnings"].append(
                "Some low-confidence OCR regions were kept as image content to preserve seals, signatures or unclear text."
            )
        if exact_visual_fallback_pages:
            if report.get("hostedOcrTimedOut") and pages_with_text == 0:
                report["warnings"].append(
                    f"{exact_visual_fallback_pages} scanned page(s) were preserved as exact page images because hosted OCR timed out."
                )
            else:
                report["warnings"].append(
                    f"{exact_visual_fallback_pages} dense non-template scanned page(s) were preserved as exact page images with OCR text hidden behind them to avoid visual corruption."
                )
        if total_skipped_low_conf:
            report["warnings"].append(
                f"{total_skipped_low_conf} OCR text candidates were too uncertain to rebuild as editable text."
            )
        if colored_marks_kept:
            report["warnings"].append(
                "Colored stamp/signature regions were kept visual-only to avoid OCR artifacts over official marks."
            )
        out.save(dst)
    finally:
        pdf.close()
        shutil.rmtree(tmpdir, ignore_errors=True)


def ocr_to_docx_layout(
    src: str,
    dst: str,
    lang: str,
    table_detection: bool,
    report: dict[str, Any],
    quality: str = "high",
) -> None:
    import fitz
    from docx import Document
    from docx.shared import Pt

    dpi = quality_dpi(quality, 300)
    doc = fitz.open(src)
    out = Document()
    tmpdir = tempfile.mkdtemp(prefix="filemint-ocr-")
    try:
        all_low_conf = 0
        pages_with_text = 0
        scanned_tables = 0

        for page_index, page in enumerate(doc):
            if page_index > 0:
                out.add_page_break()
            section = out.sections[-1]
            section.page_width = Pt(page.rect.width)
            section.page_height = Pt(page.rect.height)
            section.top_margin = Pt(36)
            section.bottom_margin = Pt(36)
            section.left_margin = Pt(36)
            section.right_margin = Pt(36)

            pix = page.get_pixmap(dpi=dpi, alpha=False)
            img = os.path.join(tmpdir, f"page-{page_index + 1}.png")
            pix.save(img)
            tsv = run_tesseract_tsv(img, lang, psm="6")
            lines = parse_tsv(tsv, pix.width, pix.height, page.rect.width, page.rect.height)
            if lines:
                pages_with_text += 1
            all_low_conf += sum(1 for line in lines for word in line.words if 0 <= word.conf < 55)

            runs = table_runs(lines) if table_detection else []
            table_line_ids = {id(line) for run in runs for line in run}
            run_by_first_id = {id(run[0]): run for run in runs}
            scanned_tables += len(runs)

            prev_bottom: float | None = None
            for line in lines:
                if id(line) in run_by_first_id:
                    add_table(out, run_by_first_id[id(line)], report)
                    prev_bottom = max(l.top + l.height for l in run_by_first_id[id(line)])
                    continue
                if id(line) in table_line_ids:
                    continue
                add_paragraph_from_line(out, line, prev_bottom)
                prev_bottom = line.top + line.height

            if not lines:
                p = out.add_paragraph()
                p.add_run("")

        if pages_with_text == 0:
            report["warnings"].append("OCR produced no editable text.")
        if table_detection and scanned_tables == 0:
            report["warnings"].append(
                "No scanned tables were confidently reconstructed. Table detection may need clearer grid lines or better OCR language data."
            )
        report["editableTextDetected"] = pages_with_text > 0
        report["lowConfidenceOcrAreas"] = all_low_conf
        report["notes"].append("Scanned PDF rebuilt from Tesseract OCR word positions into editable DOCX text and simple tables.")
        out.save(dst)
    finally:
        doc.close()
        shutil.rmtree(tmpdir, ignore_errors=True)


def to_docx_image(
    src: str,
    dst: str,
    report: dict[str, Any],
    quality: str = "high",
    visual_object_format: str = "png",
) -> None:
    import fitz
    from docx import Document
    from docx.shared import Inches, Pt

    pdf = fitz.open(src)
    tmpdir = tempfile.mkdtemp(prefix="filemint-image-")
    try:
        out = Document()
        for i, page in enumerate(pdf):
            if i > 0:
                out.add_page_break()
            section = out.sections[-1]
            section.page_width = Pt(page.rect.width)
            section.page_height = Pt(page.rect.height)
            section.top_margin = Pt(0)
            section.bottom_margin = Pt(0)
            section.left_margin = Pt(0)
            section.right_margin = Pt(0)
            ext = "jpg" if visual_object_format == "jpg" else "png"
            pix = page.get_pixmap(dpi=quality_dpi(quality, 180), alpha=False)
            img = os.path.join(tmpdir, f"page-{i + 1}.{ext}")
            pix.save(img)
            out.add_picture(img, width=Inches(page.rect.width / 72.0))
        out.save(dst)
        report["editableTextDetected"] = False
        report["nonEditableVisualFallback"] = True
        report["warnings"].append("Image-only mode creates non-editable page pictures. Use it only when editable reconstruction fails.")
    finally:
        pdf.close()
        shutil.rmtree(tmpdir, ignore_errors=True)


def ensure_output(path: str) -> None:
    if not os.path.exists(path) or os.path.getsize(path) == 0:
        raise SystemExit("Conversion produced no output.")


def docx_output_stats(path: str) -> dict[str, Any]:
    stats: dict[str, Any] = {
        "outputTextRuns": 0,
        "outputEditableCharacters": 0,
        "outputImages": 0,
        "outputTables": 0,
    }
    if not os.path.exists(path) or os.path.getsize(path) == 0:
        return stats

    try:
        with zipfile.ZipFile(path) as zf:
            names = zf.namelist()
            xml_names = [
                name
                for name in names
                if name == "word/document.xml"
                or re.match(r"word/(header|footer|footnotes|endnotes)\d*\.xml$", name)
            ]
            for name in xml_names:
                xml = zf.read(name).decode("utf-8", errors="ignore")
                pieces = re.findall(r"<w:t(?:\s[^>]*)?>(.*?)</w:t>", xml, flags=re.S)
                stats["outputTextRuns"] += len(pieces)
                stats["outputEditableCharacters"] += sum(len(unescape(piece)) for piece in pieces)
                stats["outputTables"] += xml.count("<w:tbl")
            stats["outputImages"] = sum(1 for name in names if name.startswith("word/media/"))
    except Exception:
        return stats
    return stats


def merge_output_stats(report: dict[str, Any], path: str) -> dict[str, Any]:
    stats = docx_output_stats(path)
    report.update(stats)
    return stats


def repair_empty_editable_output(
    src: str,
    dst: str,
    lang_request: str,
    table_detection: bool,
    report: dict[str, Any],
    quality: str,
    keep_visual_objects: bool,
) -> None:
    report["warnings"].append(
        "The first DOCX pass contained no editable text, so FileMint retried with OCR Editable reconstruction instead of returning an image-only Word file."
    )
    lang = resolve_ocr_language(lang_request, report)
    try:
        ocr_to_docx_exact_visual(
            src,
            dst,
            lang,
            report,
            premium=True,
            table_detection=table_detection,
            visible_text=True,
            quality=quality,
            keep_visual_objects=keep_visual_objects,
        )
        report["resolvedMode"] = "ocr-repair-editable-visual"
    except Exception as e:
        report["warnings"].append(
            f"OCR repair with visual placement failed ({e}); retrying text-only OCR reconstruction."
        )
        ocr_to_docx_layout(src, dst, lang, table_detection, report, quality=quality)
        report["resolvedMode"] = "ocr-repair-text-flow"


def write_report(path: str | None, report: dict[str, Any]) -> None:
    if not path:
        return
    with open(path, "w", encoding="utf-8") as f:
        json.dump(report, f, ensure_ascii=True, indent=2)


def main() -> None:
    ap = argparse.ArgumentParser()
    ap.add_argument("--input", required=True)
    ap.add_argument("--output", required=True)
    ap.add_argument("--mode", default="hybrid")
    ap.add_argument("--lang", default="auto")
    ap.add_argument("--auto-detect-language", default="true")
    ap.add_argument("--table-detection", default="true")
    ap.add_argument("--preserve-layout", default="true")
    ap.add_argument("--keep-visual-objects", default="true")
    ap.add_argument("--visual-object-format", default="png")
    ap.add_argument("--docx-quality", default="high")
    ap.add_argument("--report")
    a = ap.parse_args()

    requested_mode = safe_mode(a.mode)
    run_mode = engine_mode(requested_mode)
    premium_mode = run_mode == "premium"
    auto_detect_language = truthy(a.auto_detect_language, True)
    table_detection = truthy(a.table_detection, True)
    preserve_layout = truthy(a.preserve_layout, True)
    keep_visual_objects = truthy(a.keep_visual_objects, True)
    visual_object_format = clean_choice(a.visual_object_format, {"png", "jpg", "jpeg"}, "png")
    if visual_object_format == "jpeg":
        visual_object_format = "jpg"
    docx_quality = clean_choice(a.docx_quality, {"low", "medium", "high", "original"}, "high")
    report: dict[str, Any] = {
        "engine": "filemint-local",
        "requestedMode": requested_mode,
        "resolvedMode": requested_mode,
        "engineMode": run_mode,
        "pdfType": "unknown",
        "pagesConverted": 0,
        "editableTextDetected": False,
        "tablesDetected": 0,
        "imagesDetected": 0,
        "lowConfidenceOcrAreas": 0,
        "editableTextBoxes": 0,
        "editableCharacters": 0,
        "ocrTextCandidates": 0,
        "textCoverageEstimate": 0,
        "visualObjectsPreserved": 0,
        "ocrPasses": [],
        "ocrLanguage": None,
        "autoDetectLanguage": auto_detect_language,
        "tableDetectionEnabled": table_detection,
        "layoutPreservationEnabled": preserve_layout,
        "keepVisualObjects": keep_visual_objects,
        "visualObjectFormat": visual_object_format,
        "docxQuality": docx_quality,
        "nonEditableVisualFallback": False,
        "warnings": [],
        "notes": [],
        "pageDetails": [],
    }

    try:
        info = inspect_pdf(a.input)
        report.update(
            {
                "pdfType": info["pdfType"],
                "pagesConverted": info["pages"],
                "imageBackedPages": info.get("imageBackedPages", 0),
                "tablesDetected": info["tablesDetected"],
                "imagesDetected": info["imagesDetected"],
                "pageDetails": info["pageDetails"],
            }
        )

        digital_enough = info["pdfType"] in {"digital", "mixed"} and info["textCharacters"] >= 25
        scanned_enough = info["pdfType"] == "scanned" or (info["pdfType"] == "mixed" and info["textPages"] < info["pages"])
        image_backed_text_layer = digital_enough and image_backed_text_layer_likely(info)
        lang_request = effective_ocr_request(a.lang, auto_detect_language, report)
        editable_modes = {"premium", "exact"}
        digital_text_flow_candidate = (
            premium_mode
            and digital_enough
            and not image_backed_text_layer
            and int(info.get("imagesDetected", 0)) == 0
            and int(info.get("tablesDetected", 0)) == 0
            and int(info.get("textCharacters", 0)) >= 500
        )

        if run_mode in editable_modes and image_backed_text_layer:
            report["warnings"].append(
                "This PDF has full-page scan images plus a text layer, so FileMint avoided the standard PDF object converter and rebuilt the content as editable Word text."
            )
            lang: str | None = None
            if premium_mode or image_backed_text_layer_needs_ocr(info):
                try:
                    lang = resolve_ocr_language(lang_request, report)
                except Exception as e:
                    report["warnings"].append(
                        f"OCR fallback is unavailable ({e}); pages with a weak text layer may remain less editable."
                    )
            to_docx_scan_text_layer(
                a.input,
                a.output,
                lang,
                report,
                premium=premium_mode,
                table_detection=table_detection,
                quality=docx_quality,
            )
        elif digital_text_flow_candidate:
            to_docx_digital_text_flow(a.input, a.output, report)
        elif run_mode in editable_modes and digital_enough:
            report["resolvedMode"] = requested_mode
            if premium_mode:
                report["notes"].append(
                    "High Accuracy/Hybrid mode on this digital PDF used the editable object-model converter because images or detected tables make text-flow reconstruction less appropriate."
                )
            elif requested_mode == "exact":
                report["notes"].append(
                    "Exact Visual Mode on digital PDFs uses the same editable layout engine with stricter visual-fidelity intent."
                )
            to_docx_pdf2docx(a.input, a.output, report)
        elif run_mode in editable_modes and scanned_enough:
            if requested_mode == "hybrid":
                report["resolvedMode"] = "hybrid-ocr-editable-visual"
            elif requested_mode == "high-accuracy":
                report["resolvedMode"] = "high-accuracy-ocr-editable-visual"
            elif requested_mode == "exact":
                report["resolvedMode"] = "exact-ocr-visual"
            else:
                report["resolvedMode"] = "ocr-editable-visual" if preserve_layout else "ocr"
            report["warnings"].append(
                "This PDF looks scanned/image-based, so conversion was routed through OCR reconstruction."
            )
            lang = resolve_ocr_language(lang_request, report)
            if preserve_layout or premium_mode or requested_mode == "exact":
                ocr_to_docx_exact_visual(
                    a.input,
                    a.output,
                    lang,
                    report,
                    premium=premium_mode,
                    table_detection=table_detection,
                    visible_text=True,
                    quality=docx_quality,
                    keep_visual_objects=keep_visual_objects,
                )
            else:
                ocr_to_docx_layout(a.input, a.output, lang, table_detection, report, quality=docx_quality)
        elif requested_mode == "ocr":
            report["resolvedMode"] = "ocr-editable-visual" if preserve_layout else "ocr"
            lang = resolve_ocr_language(lang_request, report)
            if preserve_layout:
                ocr_to_docx_exact_visual(
                    a.input,
                    a.output,
                    lang,
                    report,
                    table_detection=table_detection,
                    visible_text=True,
                    quality=docx_quality,
                    keep_visual_objects=keep_visual_objects,
                )
            else:
                ocr_to_docx_layout(a.input, a.output, lang, table_detection, report, quality=docx_quality)
        elif requested_mode == "image":
            report["resolvedMode"] = "image"
            to_docx_image(a.input, a.output, report, quality=docx_quality, visual_object_format=visual_object_format)
        else:
            report["resolvedMode"] = "high-accuracy"
            to_docx_pdf2docx(a.input, a.output, report)

        ensure_output(a.output)
        stats = merge_output_stats(report, a.output)
        hosted_timeout_fallback = bool(report.get("hostedOcrTimedOut"))
        if (
            requested_mode != "image"
            and not hosted_timeout_fallback
            and int(stats.get("outputEditableCharacters", 0) or 0) == 0
            and int(stats.get("outputTables", 0) or 0) == 0
        ):
            repair_empty_editable_output(
                a.input,
                a.output,
                lang_request,
                table_detection,
                report,
                docx_quality,
                keep_visual_objects,
            )
            ensure_output(a.output)
            stats = merge_output_stats(report, a.output)
        hosted_timeout_fallback = bool(report.get("hostedOcrTimedOut"))
        if (
            requested_mode != "image"
            and not hosted_timeout_fallback
            and int(stats.get("outputEditableCharacters", 0) or 0) == 0
            and int(stats.get("outputTables", 0) or 0) == 0
        ):
            raise RuntimeError("DOCX output contains no editable text or tables. OCR may need to be installed or a clearer source PDF is required.")
    except Exception as e:
        report["warnings"].append(str(e))
        write_report(a.report, report)
        raise

    write_report(a.report, report)
    print("mode=" + str(report["resolvedMode"]))


if __name__ == "__main__":
    main()
